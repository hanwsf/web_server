#!/usr/bin/env python3
"""
Claude Web Server LLM - 改进版
基于标准库的多线程Web服务器，使用Claude作为后端进行语义路由
使用8085端口，支持语义分析和智能调度
"""

import http.server
import socketserver
import json
import uuid
import threading
import time
from anthropic import Anthropic  
import os
# import logging
from loguru import logger
import requests
import re
import select
import cgi
from datetime import datetime, timedelta
from urllib.parse import parse_qs, urlparse
from concurrent.futures import ThreadPoolExecutor
from pathlib import Path
import sys
import smtplib
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
from email.mime.base import MIMEBase
from email import encoders
from email.header import Header

ANTHROPIC_AUTH_TOKEN = os.getenv("ANTHROPIC_AUTH_TOKEN")                                                                                                                                  
                                                                                                                                                                                        
if not ANTHROPIC_AUTH_TOKEN:                                                                                                                                                              
    logger.warning("[AUTH] ANTHROPIC_AUTH_TOKEN 未设置，尝试从 Claude CLI 读取...")                                                                                                       
    try:                                                                                                                                                                                  
        import json                                                                                                                                                                       
        from pathlib import Path                                                                                                                                                          
        auth_file = Path.home() / ".anthropic" / "auth.json"                                                                                                                              
        if auth_file.exists():                                                                                                                                                            
            with open(auth_file, 'r') as f:                                                                                                                                               
                auth_data = json.load(f)                                                                                                                                                  
                ANTHROPIC_AUTH_TOKEN = auth_data.get("api_key")                                                                                                                           
                logger.info("[AUTH] ✅ 从 CLI 认证文件读取 Token")                                                                                                                        
    except Exception as e:                                                                                                                                                                
        logger.error(f"[AUTH] ❌ 无法读取认证: {e}")                                                                                                                                      
                                                                                                                                                                                        
# 初始化 Anthropic 客户端    todo      [CLAUDE] ❌ Claude 客户端未初始化                                                                                                                                                        
if ANTHROPIC_AUTH_TOKEN:                                                                                                                                                                  
    CLAUDE_CLIENT = Anthropic(api_key=ANTHROPIC_AUTH_TOKEN)                                                                                                                               
    logger.info("[AUTH] ✅ Claude 客户端初始化成功")                                                                                                                                      
else:                                                                                                                                                                                     
    CLAUDE_CLIENT = None                                                                                                                                                                  
    logger.error("[AUTH] ❌ 无法初始化 Claude 客户端：缺少认证信息") 


# 添加 scripts 目录到 Python 路径
SCRIPTS_DIR = Path(__file__).parent
sys.path.insert(0, str(SCRIPTS_DIR))

# 导入语义调度系统
from scheduler import SemanticScheduler, normalize_model,MODEL_MAPPING

# ============ 日志配置 ============
# logging.basicConfig(
#     level=logging.INFO,
#     format='%(asctime)s [%(levelname)-8s] %(message)s'
# )
# logger = logging.getLogger(__name__)

# ============ 配置 ============
# 尝试从.env文件加载配置
env_path = Path(__file__).parent.parent / ".env"
if env_path.exists():
    logger.info(f"从 {env_path} 加载环境变量")
    with open(env_path, 'r') as f:
        for line in f:
            line = line.strip()
            if line and not line.startswith('#'):
                try:
                    key, value = line.split('=', 1)
                    os.environ[key.strip()] = value.strip()
                except:
                    pass

WEB_PORT = int(os.environ.get('WEB_PORT', 8085))
MAX_SESSIONS = int(os.environ.get('MAX_SESSIONS', 100))  # 增加到100
SESSION_TIMEOUT = int(os.environ.get('SESSION_TIMEOUT', 3600))
WORKER_THREADS = int(os.environ.get('WORKER_THREADS', 16))  # 增加到16
CLAUDE_EXECUTABLE = os.environ.get('CLAUDE_EXECUTABLE', 'claude')
KNOWLEDGE_BASE_SCRIPT = "/home/will/.claude/skills/knowledge-base/scripts"
WORK_DIR = Path(os.environ.get('WORK_DIR', '/home/will/Downloads/opencode_p'))

# NVIDIA API配置
NVIDIA_BASE_URL = os.environ.get('NVIDIA_BASE_URL', '')
NVIDIA_API_KEY = os.environ.get('NVIDIA_API_KEY', '')
NVIDIA_MODELS = [m.strip() for m in os.environ.get('NVIDIA_MODELS', '').split(',') if m.strip()]

# Deepseek API配置
DEEPSEEK_BASE_URL = os.environ.get('DEEPSEEK_BASE_URL', '')
DEEPSEEK_API_KEY = os.environ.get('DEEPSEEK_API_KEY', '')
DEEPSEEK_MODELS = [m.strip() for m in os.environ.get('DEEPSEEK_MODELS', '').split(',') if m.strip()]

# Claude支持的模型
# CLAUDE_MODELS = [m.strip() for m in os.environ.get('CLAUDE_MODELS', 'opus,sonnet,haiku').split(',') if m.strip()]

# 超时设置
DEFAULT_TIMEOUT = int(os.environ.get('DEFAULT_TIMEOUT', 300))
NVIDIA_TIMEOUT = int(os.environ.get('NVIDIA_TIMEOUT', 120))
DEEPSEEK_TIMEOUT = int(os.environ.get('DEEPSEEK_TIMEOUT', 120))
DOCUMENT_TIMEOUT = int(os.environ.get('DOCUMENT_TIMEOUT', 900))  # 文档生成15分钟
CLAUDE_TIMEOUT = int(os.environ.get('CLAUDE_TIMEOUT', 300))

# Claude支持的模型
CLAUDE_MODELS = [m.strip() for m in os.environ.get('CLAUDE_MODELS', 'haiku,opus').split(',') if m.strip()]

KB_PATHS = {
    "KB": str(WORK_DIR / "KB"),
    "KBGX": str(WORK_DIR / "KBGX"),
    "KBW": str(WORK_DIR / "KBW")
}

logger.info(f"NVIDIA模型配置: {NVIDIA_MODELS}")
logger.info(f"Deepseek模型配置: {DEEPSEEK_MODELS}")
logger.info(f"Claude模型配置: {CLAUDE_MODELS}")
logger.info(f"Claude可执行路径: {CLAUDE_EXECUTABLE}")

# ============ 模型映射 ============
# 将前端模型选择值映射到实际的模型标识符
# 注意：映射到的模型必须是Claude通过CLI支持的模型
# MODEL_MAPPING = {
#     # GLM 模型系列
#     "glm4.7": "nvidia/z-ai/glm4.7",  # GLM4.7 映射到 Deepseek（更稳定）

#     # Deepseek 模型系列
#     "deepseek-v3.1-terminus": "nvidia/deepseek-ai/deepseek-v3.1-terminus",
#     "deepseek-v3.2": "nvidia/deepseek-ai/deepseek-v3.2",

#     # MiniMax 模型系列
#     "minimax-m2.1": "nvidia/minimaxai/minimax-m2.1",

#     # Deepseek Chat
#     "deepseek-chat": "deepseek/deepseek-chat",

# }

# def normalize_model(model: str) -> str:
#     """将前端选择的模型值标准化为支持的格式"""
#     if not model:
#         return "deepseek-v3.1-terminus"  # 使用 Deepseek 作为默认模型（更稳定）

#     # 如果有映射，使用映射值
#     if model in MODEL_MAPPING:
#         normalized = MODEL_MAPPING[model]
#         logger.info(f"[MODEL] 将前端模型 '{model}' 映射到 '{normalized}'")
#         return normalized

#     # 否则直接使用（假设前端已经传入了正确的格式）
#     logger.info(f"[MODEL] 使用前端模型 '{model}'")
#     return model

# ============ 语义调度系统 ============
# 封装模型调用器供调度器使用-这是最早的模型设置，默认模型设置
def scheduler_model_caller(prompt: str, model: str = "nvidia/minimaxai/minimax-m2.1", **kwargs) -> str:
    """供调度器使用的模型调用器 - Web 服务优化版本

    优先级顺序:
    1. 网络搜索请求 → Claude (需要WebSearch工具)
    2. Claude模型 (opus, sonnet, haiku) - 优先使用
    3. 快速降级 - 如果Claude失败，尝试降级方案
    """
    # 对于 Web 服务，优先使用Claude
    # 网络搜索请求需要特殊处理 - 必须使用 Claude 的 WebSearch 工具

    # 标准化模型参数
    model = normalize_model(model)
    logger.info(f"==model:{model}")

    timeout = kwargs.get('timeout', DEFAULT_TIMEOUT)  # 默认使用300秒，文档生成可传900秒

    # 检测网络搜索请求
    # start_marker = "用户输入: "
    # start_index = prompt.find(start_marker)
#使用dk, 这里不截留
    # if start_index != -1:
    #     # 计算用户输入开始位置
    #     input_start = start_index + len(start_marker)
    #     # 获取用户输入部分
    #     user_input = prompt[input_start:]
        
        # 取前30个字（字符）
        # prompt_lower = user_input[:30]

    # is_web_search = any([
    #     # 显式网络搜索关键词
    #     "网络搜索" in prompt_lower,
    #     "web搜索" in prompt_lower,
    #     "网络查询" in prompt_lower,
    #     "搜索网络" in prompt_lower,
    #     "互联网搜索" in prompt_lower,
    #     "在线搜索" in prompt_lower,
    #     "search the web" in prompt_lower,
    #     "web search" in prompt_lower,
    #     prompt.startswith("网络搜索"),
    #     prompt.startswith("搜索"),
    #     # 实时信息关键词 - 需要网络搜索的内容
    #     "今天" in prompt_lower or "昨天" in prompt_lower or "明天" in prompt_lower,
    #     "天气" in prompt_lower,
    #     "新闻" in prompt_lower,
    #     "股票" in prompt_lower,
    #     "汇率" in prompt_lower,
    #     "价格" in prompt_lower,
    #     "比赛" in prompt_lower or "直播" in prompt_lower,
    #     "最新" in prompt_lower,
    #     "当前" in prompt_lower or "实时" in prompt_lower,
    # ])

    # # 如果是网络搜索请求，直接使用 Claude（它有 WebSearch 工具）
    # # 注意：必须使用 Claude 模型，不能使用其他 API 模型（它们没有 WebSearch 工具）
    # if is_web_search:
    #     logger.info(f"[SCHEDULER] 检测到网络搜索请求，强制使用 Claude haiku (使用长超时), prompt:{prompt}")
    #     ws_timeout = int(os.environ.get('WEBSEARCH_TIMEOUT', '600'))
        # return call_claude(prompt, "haiku", ws_timeout)  # 强制使用 haiku 而不是用户选择的模型。

    # 优先使用用户指定的模型
    logger.info(f"[SCHEDULER] 使用用户指定的模型: {model}, timeout: {timeout}s")

    # 根据模型类型路由（传递的timeout会被使用，允许文档生成使用更长的超时）
    logger.info(f"=使用的模型：{model}")
    if model in DEEPSEEK_MODELS:
        return call_deepseek(prompt, model, timeout)
    elif model in NVIDIA_MODELS:
        return call_nvidia(prompt, model, timeout)
    elif model in CLAUDE_MODELS:
        return call_claude(prompt, model, timeout)
    else:
        # 默认使用 Claude 处理
        logger.info(f"默认使用 Claude 处理, haiku")
        return call_claude(prompt, "haiku", timeout)

# 初始化语义调度器
SKILLS_DIR = Path(os.environ.get('SKILLS_DIR', '/home/will/.claude/skills/')).expanduser()
semantic_scheduler = None  # 将在服务器启动时初始化

# ============ 全局状态 ============
SESSIONS = {}
SESSIONS_LOCK = threading.Lock()
EXECUTOR = ThreadPoolExecutor(max_workers=WORKER_THREADS)
PROGRESS_TRACKING = {}
USER_CONFIRMATIONS = {}  # 存储用户自动确认设置

# 进程跟踪 - 用于停止正在运行的子进程
RUNNING_PROCESSES = {}  # request_id -> subprocess.Popen对象
RUNNING_PROCESSES_LOCK = threading.Lock()
CURRENT_REQUEST_CONTEXT = threading.local()  # 线程本地变量，存储当前请求ID

# ============ 工具函数 ============

def run_command(cmd: list, timeout: int = 60) -> tuple:
    """运行命令并返回(成功标志, 输出)"""
    try:
        result = subprocess.run(
            cmd,
            capture_output=True,
            text=True,
            timeout=timeout
        )
        return result.returncode == 0, result.stdout + result.stderr
    except subprocess.TimeoutExpired:
        return False, "❌ 命令执行超时"
    except Exception as e:
        return False, f"❌ 错误: {str(e)}"

def call_model(prompt: str, model: str = "haiku", timeout: int = DEFAULT_TIMEOUT) -> str:
    """统一模型调用函数，根据模型类型路由到不同后端"""
    # 标准化模型参数
    model = normalize_model(model)

    # 检查是否是Deepseek模型
    if model in DEEPSEEK_MODELS:
        return call_deepseek(prompt, model, min(timeout, DEEPSEEK_TIMEOUT))

    # 检查是否是NVIDIA模型
    if model in NVIDIA_MODELS:
        return call_nvidia(prompt, model, min(timeout, NVIDIA_TIMEOUT))

    # 检查是否是Claude模型
    if model in CLAUDE_MODELS:
        return call_claude(prompt, model, min(timeout, CLAUDE_TIMEOUT))

    # 默认使用Claude处理
    logger.info(f"[MODEL] 使用 Claude 处理请求")
    return call_claude(prompt, "haiku", timeout)


def call_nvidia(prompt: str, model: str, timeout: int = NVIDIA_TIMEOUT) -> str:
    """调用NVIDIA API进行处理"""
    logger.info(f"[NVIDIA] 调用{model}: {prompt}...")
    start_time = time.time()

    if not NVIDIA_BASE_URL or not NVIDIA_API_KEY:
        return "❌ NVIDIA API未配置，请检查.env文件中的NVIDIA_BASE_URL和NVIDIA_API_KEY"

    try:
        headers = {
            "Authorization": f"Bearer {NVIDIA_API_KEY}",
            "Content-Type": "application/json"
        }

        data = {
            "model": model.replace('nvidia/',''),
            "messages": [
                {"role": "user", "content": prompt}
            ],
            "max_tokens": 8192,
            "temperature": 0.1
        }

        response = requests.post(
            f"{NVIDIA_BASE_URL}/chat/completions",
            headers=headers,
            json=data,
            timeout=timeout
        )

        elapsed = time.time() - start_time

        if response.status_code == 200:
            result = response.json()
            content = result.get('choices', [{}])[0].get('message', {}).get('content', '')
            if content:
                logger.info(f"[NVIDIA] ✅ 处理成功 ({len(content)}字, {elapsed:.1f}s)")
                return content
            else:
                logger.warning(f"[NVIDIA] ⚠️ 返回空内容 ({elapsed:.1f}s),failsafe deepseek")
                return call_deepseek(prompt, DEEPSEEK_MODELS[0], min(timeout, DEEPSEEK_TIMEOUT)) #"NVIDIA API返回空内容"
        else:
            logger.error(f"[NVIDIA] ❌ API错误: {response.status_code} ({elapsed:.1f}s)")
            return call_deepseek(prompt, DEEPSEEK_MODELS[0], min(timeout, DEEPSEEK_TIMEOUT)) #f"NVIDIA API错误: {response.status_code} - {response.text[:600]}"

    except requests.exceptions.Timeout:
        elapsed = time.time() - start_time
        logger.error(f"[NVIDIA] 超时 ({timeout}s, 实际{elapsed:.1f}s)")
        return call_deepseek(prompt, DEEPSEEK_MODELS[0], min(timeout, DEEPSEEK_TIMEOUT)) #f"⏳ NVIDIA API处理超时 ({timeout}秒)"
    except Exception as e:
        elapsed = time.time() - start_time
        logger.error(f"[NVIDIA] 错误 ({elapsed:.1f}s): {str(e)}")
        return call_deepseek(prompt, DEEPSEEK_MODELS[0], min(timeout, DEEPSEEK_TIMEOUT)) #f"❌ NVIDIA API调用失败: {str(e)}"


def call_deepseek(prompt: str, model: str, timeout: int = DEEPSEEK_TIMEOUT) -> str:
    """调用Deepseek API进行处理"""
    logger.info(f"[DEEPSEEK] 调用{model}: {prompt}...")
    start_time = time.time()

    if not DEEPSEEK_BASE_URL or not DEEPSEEK_API_KEY:
        return "❌ Deepseek API未配置，请检查.env文件中的DEEPSEEK_BASE_URL和DEEPSEEK_API_KEY"

    try:
        headers = {
            "Authorization": f"Bearer {DEEPSEEK_API_KEY}",
            "Content-Type": "application/json"
        }

        data = {
            "model": model.replace('deepseek/',''),
            "messages": [
                {"role": "user", "content": prompt}
            ],
            "max_tokens": 8192,
            "temperature": 0.1,
            "stream": False
        }

        response = requests.post(
            f"{DEEPSEEK_BASE_URL}/chat/completions",
            headers=headers,
            json=data,
            timeout=timeout
        )

        elapsed = time.time() - start_time

        if response.status_code == 200:
            result = response.json()
            content = result.get('choices', [{}])[0].get('message', {}).get('content', '')
            if content:
                logger.info(f"[DEEPSEEK] ✅ 处理成功 ({len(content)}字, {elapsed:.1f}s)")
                return content
            else:
                logger.warning(f"[DEEPSEEK] ⚠️ 返回空内容 ({elapsed:.1f}s)")
                return "Deepseek API返回空内容"
        else:
            logger.error(f"[DEEPSEEK] ❌ API错误: {response.status_code} ({elapsed:.1f}s)")
            return f"Deepseek API错误: {response.status_code} - {response.text[:800]}"

    except requests.exceptions.Timeout:
        elapsed = time.time() - start_time
        logger.error(f"[DEEPSEEK] 超时 ({timeout}s, 实际{elapsed:.1f}s)")
        return f"⏳ Deepseek API处理超时 ({timeout}秒)"
    except Exception as e:
        elapsed = time.time() - start_time
        logger.error(f"[DEEPSEEK] 错误 ({elapsed:.1f}s): {str(e)}")
        return f"❌ Deepseek API调用失败: {str(e)}"

#  现在的流程                                                                                                                                                             
                                                                                                                                                                         
#   用户输入 → 前端发送请求(202) → 后端异步处理                                                                                                                            
#       ↓                                                                                                                                                                  
#   Claude CLI接收prompt → 关闭stdin → 开始执行WebSearch                                                                                                                   
#       ↓                                                                                                                                                                  
#   前端轮询结果(每1秒) → 60秒后显示结果                                                                                                                                   
                                                                                                                                                                         
#   测试建议：                                                                                                                                                             
#   输入：查询今天上海的天气                                                                                                                                               
#   预期：60-70秒内显示天气信息（不卡死）

#  我的解决方案 ✅
# #   使用Claude CLI（用户的正常工作环境）
# def call_claude(prompt: str, model: str = "claude-haiku-4-5-20251001", timeout: int = CLAUDE_TIMEOUT, fallback_to_deepseek: bool = True, auto_confirm: bool = True) -> str:
#     """使用Claude CLI调用Claude模型

#     Args:
#         prompt: 用户输入
#         model: 模型名称
#         timeout: 超时时间(秒)
#         fallback_to_deepseek: 失败时是否降级到Deepseek
#         auto_confirm: 是否自动确认WebSearch等需要用户确认的操作
#     """
#     if model == "haiku":
#         model = "claude-haiku-4-5-20251001"

#     logger.info(f"[CLAUDE-CLI] 调用{model}: {prompt}...")
#     start_time = time.time()

#     # 获取当前请求ID（如果有）
#     current_request_id = getattr(CURRENT_REQUEST_CONTEXT, 'request_id', None)
#     process = None

#     try:
#         cmd = ["claude", "--print", "--model", model]

#         logger.debug(f"[CLAUDE-CLI] 执行命令: {' '.join(cmd)}")

#         process = subprocess.Popen(
#             cmd,
#             stdin=subprocess.PIPE,
#             stdout=subprocess.PIPE,
#             stderr=subprocess.PIPE,
#             text=True
#         )

#         # 注册进程到全局跟踪
#         if current_request_id:
#             with RUNNING_PROCESSES_LOCK:
#                 RUNNING_PROCESSES[current_request_id] = process
#                 logger.info(f"[CLAUDE-CLI] 注册进程 PID={process.pid} for request_id={current_request_id}")

#         stdout, stderr = process.communicate(input=prompt, timeout=timeout)#'Invalid API key · Please run /login\n'
#         elapsed = time.time() - start_time
#         logger.info("=======================calude stdout:{stdout}")
#         if process.returncode == 0 and stdout:
#             logger.info(f"[CLAUDE-CLI] ✅ 处理成功 ({len(stdout)}字, {elapsed:.1f}s)")
#             return stdout.strip()
#         else:
#             error_msg = stderr or stdout or "未知错误"
#             logger.error(f"[CLAUDE-CLI] 执行失败 (返回码{process.returncode}): {error_msg[:200]}")

#             # 如果CLI失败，尝试降级到Deepseek
#             if fallback_to_deepseek and DEEPSEEK_MODELS:
#                 logger.warning(f"[CLAUDE-CLI] Claude CLI因模型暂不可用失败，尝试降级到Deepseek")
#                 try:
#                     fallback_result = call_deepseek(prompt, DEEPSEEK_MODELS[0], min(timeout, DEEPSEEK_TIMEOUT))
#                     logger.info(f"[CLAUDE-CLI] ✅ Deepseek降级成功")
#                     return f"[Claude CLI执行失败，已自动使用Deepseek处理]\n\n{fallback_result}"
#                 except Exception as e:
#                     logger.error(f"[CLAUDE-CLI] Deepseek降级也失败: {str(e)}")

#             return f"❌ Claude CLI错误: {error_msg[:300]}"

#     except subprocess.TimeoutExpired:
#         elapsed = time.time() - start_time
#         logger.error(f"[CLAUDE-CLI] 超时 ({timeout}s, 实际{elapsed:.1f}s)")

#         if fallback_to_deepseek and DEEPSEEK_MODELS:
#             logger.warning(f"[CLAUDE-CLI] Claude CLI超时，尝试降级到Deepseek")
#             try:
#                 fallback_result = call_deepseek(prompt, DEEPSEEK_MODELS[0], min(timeout, DEEPSEEK_TIMEOUT))
#                 logger.info(f"[CLAUDE-CLI] ✅ Deepseek降级成功")
#                 return f"[Claude CLI处理超时，已自动使用Deepseek处理]\n\n{fallback_result}"
#             except Exception as e:
#                 logger.error(f"[CLAUDE-CLI] Deepseek降级也失败: {str(e)}")

#         return f"⏳ Claude CLI处理超时 ({timeout}秒)"

#     except Exception as e:
#         elapsed = time.time() - start_time
#         logger.error(f"[CLAUDE-CLI] 错误 ({elapsed:.1f}s): {str(e)}")

#         if fallback_to_deepseek and DEEPSEEK_MODELS:
#             logger.warning(f"[CLAUDE-CLI] Claude CLI出错，尝试降级到Deepseek")
#             try:
#                 fallback_result = call_deepseek(prompt, DEEPSEEK_MODELS[0], min(timeout, DEEPSEEK_TIMEOUT))
#                 logger.info(f"[CLAUDE-CLI] ✅ Deepseek降级成功")
#                 return f"[Claude CLI出错，已自动使用Deepseek处理]\n\n{fallback_result}"
#             except Exception as e2:
#                 logger.error(f"[CLAUDE-CLI] Deepseek降级也失败: {str(e2)}")

#         return f"❌ Claude CLI错误: {str(e)}"

#     finally:
#         # 清理进程跟踪
#         if current_request_id:
#             with RUNNING_PROCESSES_LOCK:
#                 if current_request_id in RUNNING_PROCESSES:
#                     del RUNNING_PROCESSES[current_request_id]
#                     logger.info(f"[CLAUDE-CLI] 清理进程跟踪 request_id={current_request_id}")
def call_claude(prompt: str, model: str = "claude-haiku-4-5-20251001",                                                                                                                    
                  timeout: int = CLAUDE_TIMEOUT, fallback_to_deepseek: bool = True,                                                                                                         
                  auto_confirm: bool = True) -> str:                                                                                                                                        
    """                                                                                                                                                                                   
    使用 Anthropic SDK 调用 Claude（使用 ANTHROPIC_AUTH_TOKEN）                                                                                                                           
                                                                                                                                                                                        
    【改进】：                                                                                                                                                                            
    - ✅ 无需 subprocess，直接 HTTP 调用                                                                                                                                                  
    - ✅ 并发处理能力强（不占用进程）                                                                                                                                                     
    - ✅ 启动快速（无进程启动延迟）                                                                                                                                                       
    - ✅ 更稳定（无 CLI 认证问题）                                                                                                                                                        
    """                                                                                                                                                                                   
                                                                                                                                                                                        
    if not CLAUDE_CLIENT:                                                                                                                                                                 
        logger.error("[CLAUDE] ❌ Claude 客户端未初始化")                                                                                                                                 
        if fallback_to_deepseek and DEEPSEEK_MODELS:                                                                                                                                      
            return call_deepseek(prompt, DEEPSEEK_MODELS[0], timeout)                                                                                                       
        else:                                                                                                                                                                             
            raise RuntimeError("Claude 客户端未初始化且无可用的 fallback")                                                                                                                
                                                                                                                                                                                        
    if model == "haiku":                                                                                                                                                                  
        model = "claude-haiku-4-5-20251001"                                                                                                                                               
                                                                                                                                                                                        
    logger.info(f"[CLAUDE] 📞 调用 {model}")                                                                                                                                              
    logger.debug(f"[CLAUDE] Prompt: {prompt[:100]}...")                                                                                                                                   
                                                                                                                                                                                        
    start_time = time.time()                                                                                                                                                              
    current_request_id = getattr(CURRENT_REQUEST_CONTEXT, 'request_id', None)                                                                                                             
                                                                                                                                                                                        
    try:                                                                                                                                                                                  
        # 【关键】直接调用 SDK，使用 ANTHROPIC_AUTH_TOKEN                                                                                                                                 
        message = CLAUDE_CLIENT.messages.create(                                                                                                                                          
            model=model,                                                                                                                                                                  
            max_tokens=4096,                                                                                                                                                              
            messages=[                                                                                                                                                                    
                {"role": "user", "content": prompt}                                                                                                                                       
            ],                                                                                                                                                                            
            timeout=timeout                                                                                                                                                               
        )                                                                                                                                                                                 
                                                                                                                                                                                        
        elapsed = time.time() - start_time                                                                                                                                                
        result = message.content[0].text                                                                                                                                                  
                                                                                                                                                                                        
        logger.info(f"[CLAUDE] ✅ 成功 (耗时 {elapsed:.1f}s, {len(result)} 字符)")                                                                                                        
                                                                                                                                                                                        
        # 记录请求                                                                                                                                                                        
        if current_request_id:                                                                                                                                                            
            logger.debug(f"[CLAUDE] request_id={current_request_id}, model={model}, elapsed={elapsed:.1f}s")                                                                              
                                                                                                                                                                                        
        return result                                                                                                                                                                     
                                                                                                                                                                                        
    except TimeoutError:                                                                                                                                                                  
        elapsed = time.time() - start_time                                                                                                                                                
        logger.error(f"[CLAUDE] ⏱️ 超时 ({elapsed:.1f}s)")                                                                                                                                
                                                                                                                                                                                        
        if fallback_to_deepseek and DEEPSEEK_MODELS:                                                                                                                                      
            logger.info("[CLAUDE] 降级到 DeepSeek...")                                                                                                                                    
            return call_deepseek(prompt, DEEPSEEK_MODELS[0], timeout, auto_confirm)                                                                                                       
        else:                                                                                                                                                                             
            raise                                                                                                                                                                         
                                                                                                                                                                                        
    except Exception as e:                                                                                                                                                                
        elapsed = time.time() - start_time                                                                                                                                                
        error_msg = str(e)                                                                                                                                                                
                                                                                                                                                                                        
        logger.error(f"[CLAUDE] ❌ 错误 ({elapsed:.1f}s): {error_msg}")                                                                                                                   
                                                                                                                                                                                        
        # 检查是否是认证错误                                                                                                                                                              
        if "401" in error_msg or "authentication" in error_msg.lower():                                                                                                                   
            logger.error("[CLAUDE] 认证错误 - 检查 ANTHROPIC_AUTH_TOKEN 是否有效")                                                                                                        
                                                                                                                                                                                        
        # 降级到 DeepSeek                                                                                                                                                                 
        if fallback_to_deepseek and DEEPSEEK_MODELS:                                                                                                                                      
            logger.info("[CLAUDE] 自动降级到 DeepSeek")                                                                                                                                   
            try:                                                                                                                                                                          
                return call_deepseek(prompt, DEEPSEEK_MODELS[0], min(timeout, DEEPSEEK_TIMEOUT))                                                                                                   
            except Exception as fallback_error:                                                                                                                                           
                logger.error(f"[CLAUDE] DeepSeek 也失败了: {fallback_error}")                                                                                                             
                raise                                                                                                                                                                     
        else:                                                                                                                                                                             
            raise           
def list_kb_files(kb_name: str = "KB") -> str:
    """列出知识库文件"""
    if kb_name not in KB_PATHS:
        return f"❌ 知识库'{kb_name}'不存在"

    logger.info(f"[KB] 列出{kb_name}文件...")

    cmd = [
        "python3",
        f"{KNOWLEDGE_BASE_SCRIPT}/list_documents.py",
        "--kb-path", kb_name,
        "--db-path", str(WORK_DIR / ".knowledge_base")
    ]

    success, output = run_command(cmd, timeout=300)

    if success:
        logger.info(f"[KB] ✅ 列出成功")
        return output
    else:
        return f"❌ 列出失败: {output}"

def search_kb(query: str, kb_name: str = "KB") -> str:
    """搜索知识库"""
    if kb_name not in KB_PATHS:
        return f"❌ 知识库'{kb_name}'不存在"

    logger.info(f"[KB] 搜索{kb_name}: {query[:80]}...")

    # 修复：使用正确的数据库路径格式
    db_path = str(WORK_DIR / ".knowledge_base" / f"{kb_name.lower()}_index.json")

    cmd = [
        "python3",
        f"{KNOWLEDGE_BASE_SCRIPT}/search_knowledge_base.py",
        query,
        "--kb-path", kb_name,
        "--db-path", db_path,
        "--format", "answer"
    ]

    success, output = run_command(cmd, timeout=300)

    if success:
        logger.info(f"[KB] ✅ 搜索成功")
        return output
    else:
        return f"❌ 搜索失败: {output}"

def _fetch_with_requests(url: str) -> tuple:
    """使用 requests 获取网页内容（快速但不支持JS）"""
    headers = {
        'User-Agent': 'Mozilla/5.0 (X11; Linux x86_64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36',
        'Accept': 'text/html,application/xhtml+xml,application/xml;q=0.9,image/webp,*/*;q=0.8',
        'Accept-Language': 'zh-CN,zh;q=0.9,en;q=0.8',
    }
    resp = requests.get(url, headers=headers, timeout=30)
    resp.raise_for_status()
    resp.encoding = resp.apparent_encoding or 'utf-8'
    return resp.text, "requests"


def _fetch_with_playwright(url: str) -> tuple:
    """使用 Playwright 获取网页内容（支持JS动态渲染）"""
    from playwright.sync_api import sync_playwright

    logger.info(f"[KB] 使用 Playwright 渲染页面: {url}")

    with sync_playwright() as p:
        browser = p.chromium.launch(headless=True)
        context = browser.new_context(
            user_agent='Mozilla/5.0 (X11; Linux x86_64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36',
            viewport={'width': 1920, 'height': 1080}
        )
        page = context.new_page()

        try:
            page.goto(url, wait_until='networkidle', timeout=60000)
            # 等待页面加载完成
            page.wait_for_timeout(2000)
            html = page.content()
        finally:
            browser.close()

    return html, "playwright"


def _is_valid_content(text: str, title: str) -> bool:
    """检查提取的内容是否有效"""
    # 检查是否是反爬虫页面
    invalid_patterns = [
        '安全验证', '验证码', 'captcha', 'security check',
        '请完成验证', '人机验证', 'robot check',
        '访问过于频繁', '请稍后再试'
    ]
    title_lower = title.lower()
    text_lower = text.lower()

    for pattern in invalid_patterns:
        if pattern in title_lower or pattern in text_lower:
            return False

    # 内容太短可能是空页面或错误页面
    if len(text) < 200:
        return False

    return True


def extract_url_to_docx(url: str, kb_name: str) -> dict:
    """从URL提取内容并保存为DOCX，然后自动索引

    策略：
    1. 首先尝试 requests（快速）
    2. 如果内容无效，使用 Playwright（支持JS渲染）
    """
    from bs4 import BeautifulSoup
    from docx import Document

    logger.info(f"[KB] 提取URL内容: {url}")

    html_content = None
    fetch_method = None

    # 第一步：尝试 requests
    try:
        html_content, fetch_method = _fetch_with_requests(url)
    except Exception as e:
        logger.warning(f"[KB] requests 获取失败: {e}")

    # 解析并检查内容
    if html_content:
        soup = BeautifulSoup(html_content, 'html.parser')
        title = soup.title.string.strip() if soup.title and soup.title.string else urlparse(url).netloc

        # 移除脚本和样式
        for tag in soup(["script", "style", "nav", "footer", "header", "aside", "noscript"]):
            tag.decompose()

        article = soup.find('article') or soup.find('main') or soup.find('body')
        text = article.get_text(separator='\n', strip=True) if article else ''

        # 检查内容是否有效
        if not _is_valid_content(text, title):
            logger.info(f"[KB] requests 获取的内容无效，切换到 Playwright")
            html_content = None

    # 第二步：如果 requests 失败或内容无效，使用 Playwright
    if not html_content:
        try:
            html_content, fetch_method = _fetch_with_playwright(url)
        except Exception as e:
            logger.error(f"[KB] Playwright 获取失败: {e}")
            return {"success": False, "error": f"无法获取网页内容: {str(e)}"}

    # 重新解析 Playwright 获取的内容
    soup = BeautifulSoup(html_content, 'html.parser')

    # 提取标题
    title = soup.title.string.strip() if soup.title and soup.title.string else urlparse(url).netloc
    title = re.sub(r'[\\/:*?"<>|]', '_', title)[:100]

    # 移除脚本和样式
    for tag in soup(["script", "style", "nav", "footer", "header", "aside", "noscript", "iframe"]):
        tag.decompose()

    # 获取主要内容
    article = soup.find('article') or soup.find('main') or soup.find('div', class_=re.compile(r'content|article|post|entry')) or soup.find('body')
    text = article.get_text(separator='\n', strip=True) if article else ''

    # 清理文本：去除多余空行
    lines = [line.strip() for line in text.split('\n') if line.strip()]
    text = '\n'.join(lines)

    if not text or len(text) < 100:
        return {"success": False, "error": "无法提取网页内容（页面可能需要登录或内容为空）"}

    # 创建DOCX文档
    doc = Document()
    doc.add_heading(title, 0)
    doc.add_paragraph(f"来源: {url}")
    doc.add_paragraph(f"提取方式: {fetch_method}")
    doc.add_paragraph("")

    for para in lines:
        if para:
            doc.add_paragraph(para)

    # 保存文件
    kb_path = KB_PATHS[kb_name]
    filename = f"{title}.docx"
    filepath = Path(kb_path) / filename

    counter = 1
    while filepath.exists():
        filename = f"{title}_{counter}.docx"
        filepath = Path(kb_path) / filename
        counter += 1

    doc.save(str(filepath))
    logger.info(f"[KB] 已保存: {filepath} (使用 {fetch_method})")

    # 自动索引
    index_result = run_index_for_kb(kb_name)

    return {
        "success": True,
        "message": f"URL内容已提取并保存为 {filename}",
        "filename": filename,
        "method": fetch_method,
        "indexed": index_result
    }


def save_uploaded_file(file_data: bytes, filename: str, kb_name: str) -> dict:
    """保存上传的文件并自动索引"""
    kb_path = KB_PATHS[kb_name]
    filepath = Path(kb_path) / filename

    # 避免重名
    counter = 1
    base_name = filepath.stem
    suffix = filepath.suffix
    while filepath.exists():
        filename = f"{base_name}_{counter}{suffix}"
        filepath = Path(kb_path) / filename
        counter += 1

    # 保存文件
    with open(filepath, 'wb') as f:
        f.write(file_data)

    logger.info(f"[KB] 已上传: {filepath}")

    # 自动索引
    index_result = run_index_for_kb(kb_name)

    return {
        "success": True,
        "message": f"文档 {filename} 已上传",
        "filename": filename,
        "indexed": index_result
    }


def run_index_for_kb(kb_name: str) -> bool:
    """运行知识库索引"""
    logger.info(f"[KB] 开始索引 {kb_name}...")

    cmd = [
        "python3",
        f"{KNOWLEDGE_BASE_SCRIPT}/index_documents.py",
        "--kb-path", KB_PATHS[kb_name],
        "--db-path", str(WORK_DIR / ".knowledge_base")
    ]

    success, output = run_command(cmd, timeout=300)

    if success:
        logger.info(f"[KB] ✅ 索引完成")
    else:
        logger.error(f"[KB] ❌ 索引失败: {output}")

    return success


def read_file(filepath: str) -> str:
    """读取文件内容"""
    try:
        with open(filepath, 'r', encoding='utf-8') as f:
            content = f.read()
        logger.info(f"[FILE] 读取成功: {filepath}")
        return content
    except Exception as e:
        logger.error(f"[FILE] 读取失败: {str(e)}")
        return f"❌ 无法读取文件: {str(e)}"

# ============ 邮件发送功能 ============

def send_email(to_email: str, subject: str, body: str, attachment_path: str = None) -> bool:
    """发送邮件，支持附件"""
    try:
        # 从环境变量读取邮件配置
        smtp_host = os.getenv('EMAIL_HOST', 'smtp.163.com')
        smtp_port = int(os.getenv('EMAIL_PORT', 465))
        sender_email = os.getenv('EMAIL_HOST_USER', '')
        sender_password = os.getenv('EMAIL_HOST_PASSWORD', '')
        use_tls = os.getenv('EMAIL_USE_TLS', 'True').lower() == 'true'

        if not sender_email or not sender_password:
            logger.error("[Email] ❌ 邮件配置不完整（EMAIL_HOST_USER 或 PASSWORD 未设置）")
            return False

        # 创建邮件
        msg = MIMEMultipart()
        msg['From'] = sender_email
        msg['To'] = to_email
        msg['Subject'] = subject

        # 添加邮件正文
        msg.attach(MIMEText(body, 'plain', 'utf-8'))

        # 添加附件
        if attachment_path and os.path.exists(attachment_path):
            try:
                filename = os.path.basename(attachment_path)
                # 根据文件扩展名确定 MIME 类型
                if filename.lower().endswith('.docx'):
                    maintype = 'application'
                    subtype = 'vnd.openxmlformats-officedocument.wordprocessingml.document'
                elif filename.lower().endswith('.pptx'):
                    maintype = 'application'
                    subtype = 'vnd.openxmlformats-officedocument.presentationml.presentation'
                elif filename.lower().endswith('.xlsx'):
                    maintype = 'application'
                    subtype = 'vnd.openxmlformats-officedocument.spreadsheetml.sheet'
                elif filename.lower().endswith('.pdf'):
                    maintype = 'application'
                    subtype = 'pdf'
                else:
                    maintype = 'application'
                    subtype = 'octet-stream'

                with open(attachment_path, 'rb') as attachment:
                    part = MIMEBase(maintype, subtype)
                    part.set_payload(attachment.read())
                    encoders.encode_base64(part)

                    # 正确处理中文文件名和特殊字符
                    # 使用 RFC 2231 编码以支持中文文件名
                    try:
                        # 尝试用 ASCII 编码，如果失败则使用 UTF-8 编码
                        filename.encode('ascii')
                        filename_header = f'attachment; filename="{filename}"'
                    except UnicodeEncodeError:
                        # 包含非 ASCII 字符，使用 RFC 2231 编码
                        from urllib.parse import quote
                        encoded_filename = quote(filename, safe='')
                        filename_header = f'attachment; filename*=utf-8\'\'{encoded_filename}'

                    part.add_header('Content-Disposition', filename_header)
                    msg.attach(part)
                logger.info(f"[Email] 已添加附件: {attachment_path} (MIME: {maintype}/{subtype})")
            except Exception as e:
                logger.error(f"[Email] ❌ 附件添加失败: {e}")
                return False

        # 发送邮件
        try:
            if use_tls:
                # 使用 SSL/TLS（通常是端口 465）
                with smtplib.SMTP_SSL(smtp_host, smtp_port, timeout=10) as server:
                    server.login(sender_email, sender_password)
                    server.send_message(msg)
            else:
                # 使用 STARTTLS（通常是端口 587）
                with smtplib.SMTP(smtp_host, smtp_port, timeout=10) as server:
                    server.starttls()
                    server.login(sender_email, sender_password)
                    server.send_message(msg)

            logger.info(f"[Email] ✅ 邮件发送成功: {to_email}")
            return True
        except smtplib.SMTPAuthenticationError:
            logger.error("[Email] ❌ 邮件认证失败（用户名或密码错误）")
            return False
        except smtplib.SMTPException as e:
            logger.error(f"[Email] ❌ SMTP 错误: {e}")
            return False

    except Exception as e:
        logger.error(f"[Email] ❌ 邮件发送异常: {e}")
        return False

def apply_text_formatting(paragraph, text: str):
    """
    将 Markdown 格式的文本应用到段落中
    支持: **bold** *italic* ***bold italic*** `code`
    """
    import re

    if not text:
        return

    # 首先清除段落中已有的文本
    for run in paragraph.runs:
        run._element.getparent().remove(run._element)

    # 使用特殊标记替换 Markdown 格式
    # 首先处理 ***bold_italic***（最高优先级）
    text = re.sub(r'\*\*\*([^*]+?)\*\*\*', r'___BOLD_ITALIC_START___\1___BOLD_ITALIC_END___', text)
    # 然后处理 **bold**
    text = re.sub(r'\*\*([^*]+?)\*\*', r'___BOLD_START___\1___BOLD_END___', text)
    # 然后处理 *italic*
    text = re.sub(r'\*([^*\n]+?)\*', r'___ITALIC_START___\1___ITALIC_END___', text)
    # 最后处理 `code`
    text = re.sub(r'`([^`]+?)`', r'___CODE_START___\1___CODE_END___', text)

    # 现在解析处理后的文本，并按格式添加到段落
    current_format = None
    current_text = ""
    i = 0
    processed_text = text

    while i < len(processed_text):
        remaining = processed_text[i:]

        if remaining.startswith('___BOLD_ITALIC_START___'):
            # 先输出当前积累的文本
            if current_text:
                run = paragraph.add_run(current_text)
                if current_format == 'bold':
                    run.bold = True
                elif current_format == 'italic':
                    run.italic = True
                elif current_format == 'bold_italic':
                    run.bold = True
                    run.italic = True
                elif current_format == 'code':
                    run.font.name = 'Courier New'
                    run.font.size = 10
                current_text = ""

            i += len('___BOLD_ITALIC_START___')
            current_format = 'bold_italic'

        elif remaining.startswith('___BOLD_ITALIC_END___'):
            if current_text:
                run = paragraph.add_run(current_text)
                run.bold = True
                run.italic = True
                current_text = ""
            i += len('___BOLD_ITALIC_END___')
            current_format = None

        elif remaining.startswith('___BOLD_START___'):
            if current_text:
                run = paragraph.add_run(current_text)
                if current_format == 'bold':
                    run.bold = True
                elif current_format == 'italic':
                    run.italic = True
                elif current_format == 'code':
                    run.font.name = 'Courier New'
                    run.font.size = 10
                current_text = ""

            i += len('___BOLD_START___')
            current_format = 'bold'

        elif remaining.startswith('___BOLD_END___'):
            if current_text:
                run = paragraph.add_run(current_text)
                run.bold = True
                current_text = ""
            i += len('___BOLD_END___')
            current_format = None

        elif remaining.startswith('___ITALIC_START___'):
            if current_text:
                run = paragraph.add_run(current_text)
                if current_format == 'bold':
                    run.bold = True
                elif current_format == 'code':
                    run.font.name = 'Courier New'
                    run.font.size = 10
                current_text = ""

            i += len('___ITALIC_START___')
            current_format = 'italic'

        elif remaining.startswith('___ITALIC_END___'):
            if current_text:
                run = paragraph.add_run(current_text)
                run.italic = True
                current_text = ""
            i += len('___ITALIC_END___')
            current_format = None

        elif remaining.startswith('___CODE_START___'):
            if current_text:
                run = paragraph.add_run(current_text)
                if current_format == 'bold':
                    run.bold = True
                elif current_format == 'italic':
                    run.italic = True
                current_text = ""

            i += len('___CODE_START___')
            current_format = 'code'

        elif remaining.startswith('___CODE_END___'):
            if current_text:
                run = paragraph.add_run(current_text)
                run.font.name = 'Courier New'
                run.font.size = 10
                current_text = ""
            i += len('___CODE_END___')
            current_format = None

        else:
            current_text += processed_text[i]
            i += 1

    # 输出最后剩余的文本
    if current_text:
        run = paragraph.add_run(current_text)
        if current_format == 'bold':
            run.bold = True
        elif current_format == 'italic':
            run.italic = True
        elif current_format == 'bold_italic':
            run.bold = True
            run.italic = True
        elif current_format == 'code':
            run.font.name = 'Courier New'
            run.font.size = 10


def content_to_docx(content: str, title: str = "文档") -> str:
    """
    将 Markdown 格式内容转换为 DOCX 文件
    支持的格式:
    - # 标题, ## 二级标题, ### 三级标题, #### 四级标题
    - **粗体** *斜体* ***粗体斜体*** `代码`
    - - 无序列表项
    - 1. 有序列表项
    """
    try:
        from docx import Document
        import re

        doc = Document()

        # 添加标题
        title_heading = doc.add_heading(title, 0)
        doc.add_paragraph("")

        # 处理 Markdown 格式
        lines = content.split('\n')
        i = 0
        while i < len(lines):
            line = lines[i]

            # 处理标题
            if line.startswith('#### '):
                heading = doc.add_heading('', level=4)
                apply_text_formatting(heading, line[5:])
            elif line.startswith('### '):
                heading = doc.add_heading('', level=3)
                apply_text_formatting(heading, line[4:])
            elif line.startswith('## '):
                heading = doc.add_heading('', level=2)
                apply_text_formatting(heading, line[3:])
            elif line.startswith('# '):
                heading = doc.add_heading('', level=1)
                apply_text_formatting(heading, line[2:])
            # 处理无序列表
            elif line.startswith('- ') or line.startswith('* '):
                p = doc.add_paragraph('', style='List Bullet')
                apply_text_formatting(p, line[2:])
            # 处理有序列表
            elif re.match(r'^\d+\.\s', line):
                # 提取序号后的内容
                content_text = re.sub(r'^\d+\.\s', '', line)
                p = doc.add_paragraph('', style='List Number')
                apply_text_formatting(p, content_text)
            # 处理普通段落
            elif line.strip():
                p = doc.add_paragraph()
                apply_text_formatting(p, line.strip())
            else:
                # 空行
                doc.add_paragraph("")

            i += 1

        # 保存文件
        filename = f"/tmp/{title}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        doc.save(filename)
        logger.info(f"[DOCX] 已生成文档: {filename}")
        return filename

    except Exception as e:
        logger.error(f"[DOCX] 生成失败: {e}")
        return None

# ============ Session类 ============

class Session:
    def __init__(self, session_id: str):
        self.session_id = session_id
        self.created_at = datetime.now()
        self.last_activity = datetime.now()
        self.messages = []
        self.lock = threading.Lock()

    def add_message(self, role: str, content: str):
        """添加消息"""
        with self.lock:
            self.messages.append({
                "role": role,
                "content": content,
                "timestamp": datetime.now().isoformat()
            })
            self.last_activity = datetime.now()

    def get_messages(self):
        """获取消息"""
        with self.lock:
            return self.messages.copy()

# ============ HTTP处理器 ============

class WebHandler(http.server.BaseHTTPRequestHandler):
    """HTTP请求处理器"""

    def log_message(self, format, *args):
        """自定义日志输出"""
        logger.info(format % args)

    def do_GET(self):
        """处理GET请求"""
        parsed = urlparse(self.path)

        if parsed.path == "/":
            self.send_html_response(self.get_frontend_html())

        elif parsed.path == "/api/status":
            with SESSIONS_LOCK:
                active = sum(1 for s in SESSIONS.values())
                self.send_json_response(200, {
                    "status": "running",
                    "sessions": len(SESSIONS),
                    "timestamp": datetime.now().isoformat()
                })

        elif parsed.path == "/api/skills":
            self._handle_list_skills({})

        elif parsed.path == "/api/agents":
            self._handle_list_agents({})

        elif parsed.path == "/api/skills/summary":
            self._handle_skills_summary({})

        elif parsed.path.startswith("/api/result/"):
            self._handle_get_result()

        elif parsed.path.startswith("/api/download/"):
            self._handle_download_file(parsed.path)

        else:
            self.send_json_response(404, {"error": "Not found"})

    def _submit_async_task(self, handler_func, data, response_immediately=True):
        """提交异步任务到线程池

        Args:
            handler_func: 处理函数
            data: 请求数据
            response_immediately: 是否立即返回202响应（不等待完成）

        Returns:
            request_id: 请求追踪ID
        """
        request_id = str(uuid.uuid4())

        # 记录请求开始
        PROGRESS_TRACKING[request_id] = {
            "status": "running",
            "start_time": datetime.now().isoformat(),
            "result": None,
            "error": None
        }

        def task_wrapper():
            """任务包装器 - 执行处理函数并记录结果"""
            try:
                logger.info(f"[ASYNC] 开始处理请求 {request_id}")

                # 设置线程本地的request_id，供call_claude等函数使用
                CURRENT_REQUEST_CONTEXT.request_id = request_id

                # 临时保存原始的 send_json_response 方法
                original_send = self.send_json_response
                captured_response = {"status_code": 200, "data": None}

                def capture_response(status_code, data):
                    """捕获响应而不是发送给客户端"""
                    captured_response["status_code"] = status_code
                    captured_response["data"] = data
                    logger.info(f"[ASYNC] 捕获响应 {request_id}: {status_code}")

                # 临时替换 send_json_response
                self.send_json_response = capture_response

                try:
                    # 执行处理函数
                    handler_func(data)
                finally:
                    # 恢复原始方法
                    self.send_json_response = original_send

                # 存储结果
                PROGRESS_TRACKING[request_id]["status"] = "completed"
                PROGRESS_TRACKING[request_id]["result"] = {
                    "status_code": captured_response["status_code"],
                    "data": captured_response["data"]
                }
                PROGRESS_TRACKING[request_id]["end_time"] = datetime.now().isoformat()
                logger.info(f"[ASYNC] 完成请求 {request_id}")

            except Exception as e:
                logger.error(f"[ASYNC] 请求 {request_id} 失败: {e}", exc_info=True)
                PROGRESS_TRACKING[request_id]["status"] = "failed"
                PROGRESS_TRACKING[request_id]["error"] = str(e)
                PROGRESS_TRACKING[request_id]["end_time"] = datetime.now().isoformat()

        # 重要：先发送202响应，再提交任务到线程池
        # 如果顺序反过来，task_wrapper可能在线程池中先启动并替换send_json_response，
        # 导致202响应被capture_response捕获而不是发送给客户端（race condition）
        if response_immediately:
            # 立即返回202 Accepted
            self.send_json_response(202, {
                "request_id": request_id,
                "status": "accepted",
                "message": "请求已接收，正在处理中",
                "result_url": f"/api/result/{request_id}"
            })

        # 提交到线程池（在发送202响应之后）
        EXECUTOR.submit(task_wrapper)

        return request_id
# 前端所有用户消息都通过 /api/claude 这个统一入口发送：

#   流程说明：

#   1. 前端 (server.py)
#   fetch('/api/claude', {
#       method: 'POST',
#       headers: { 'Content-Type': 'application/json' },
#       body: JSON.stringify({ prompt: message, model, kb })
#   })
#   - 用户输入任何消息都会发送到 /api/claude
#   - 携带参数：prompt（用户消息）、model（选择的模型）、kb（选择的知识库）

#   2. 后端路由
#   elif parsed.path == "/api/claude":
#       # Claude调用：异步处理（最可能耗时）
#       self._submit_async_task(self._handle_claude_call, data)

#   3. 语义调度处理
#   if use_semantic_scheduler and semantic_scheduler is not None:
#       # 使用 semantic_scheduler.process 来分析意图并路由
#       result = semantic_scheduler.process(prompt, {"kb": frontend_kb, "model": model})

#   为什么要这样设计？
#   - 前端简化：不需要根据不同操作调用不同 API 端点
#   - 智能路由：后端通过语义调度系统（semantic_scheduler）自动判断用户的意图
#   → list_docs（列出文件）
#   → search（搜索知识库）
#   → summarize_doc（总结）
#   → web_search（网络搜索）
#   → direct（直接用 LLM 回答）

#   这样用户只需要输入自然语言，系统自动识别意图并分发给相应的处理流程。  
    def do_POST(self):
        """处理POST请求"""
        parsed = urlparse(self.path)
        print(f"parsed:{parsed}")

        # 文件上传需要特殊处理（multipart/form-data），不能先读取 body 为 JSON
        if parsed.path == "/api/kb/upload-file":
            self._handle_upload_file()
            return

        content_length = int(self.headers.get('Content-Length', 0))
        body = self.rfile.read(content_length).decode('utf-8')

        try:
            data = json.loads(body) if body else {}
        except:
            data = {}

        if parsed.path == "/api/session":
            self._handle_create_session(data)

        elif parsed.path == "/api/send":
            self._handle_send_message(data)

        elif parsed.path.startswith("/api/result/"):
            self._handle_get_result()

        elif parsed.path == "/api/kb/list":
            # KB列表操作：直接处理（不使用异步）
            self._handle_list_kb(data)

        elif parsed.path == "/api/kb/search":
            # KB搜索：直接处理（不使用异步）
            self._handle_search_kb(data)

        elif parsed.path == "/api/kb/upload-url":
            # URL内容提取并保存为DOCX
            self._handle_upload_url(data)

        elif parsed.path == "/api/file/read":
            self._handle_read_file(data)

        elif parsed.path == "/api/email/send-content":
            # 发送文本内容到邮箱  
            self._handle_send_content_email(data)

        elif parsed.path == "/api/email/send-file":
            # 发送文件到邮箱
            self._handle_send_file_email(data)

        elif parsed.path == "/api/email/send-kb":
            # 发送知识库文件到邮箱
            self._handle_send_kb_email(data)

        elif parsed.path == "/api/email/send-article":
            # 发送文章到邮箱（支持 Markdown 转 DOCX）
            self._handle_send_article_email(data)
# 前端所有用户消息都通过 /api/claude 这个统一入口发送
# curl -X POST http://localhost:8085/api/claude \
# -H "Content-Type: application/json" \
# -d '{"prompt": "写一篇关于AI的短文，然后发送到 hanwsf@163.com"}'
# {"request_id": "c4dcb763-a9d0-4ddf-9185-8b9ac5e5adc2", "status": "accepted", "message": "\u8bf7\u6c42\u5df2\u63a5\u6536\uff0c\u6b63\u5728\u5904\u7406\u4e2d", "result_url": "/api/result/c4dcb763-a9d0-4ddf-9185-8b9ac5e5adc2"}
        elif parsed.path == "/api/claude":
            # Claude调用：异步处理（最可能耗时）
            self._submit_async_task(self._handle_claude_call, data)

        elif parsed.path == "/api/web-search":
            # 网络搜索：异步处理（需要60秒左右执行Web搜索）
            self._submit_async_task(self._handle_web_search_dk, data)
            # self._submit_async_task(self._handle_web_search, data)

        elif parsed.path == "/api/web-search-dk":
            # 网络搜索（使用dk-search服务）：异步处理
            self._submit_async_task(self._handle_web_search_dk, data)

        elif parsed.path == "/api/skills":
            self._handle_list_skills(data)

        elif parsed.path == "/api/agents":
            self._handle_list_agents(data)

        elif parsed.path == "/api/skills/summary":
            self._handle_skills_summary(data)

        elif parsed.path == "/api/stop":
            self._handle_stop_process(data)

        else:
            self.send_json_response(404, {"error": "Endpoint not found"})

    def _handle_stop_process(self, data):
        """停止正在运行的进程"""
        request_id = data.get('request_id')

        if not request_id:
            self.send_json_response(400, {"error": "Missing request_id"})
            return

        logger.info(f"[STOP] 收到停止请求: request_id={request_id}")

        process_killed = False
        with RUNNING_PROCESSES_LOCK:
            if request_id in RUNNING_PROCESSES:
                process = RUNNING_PROCESSES[request_id]
                try:
                    # 终止进程及其子进程
                    import signal
                    import os

                    # 获取进程组ID并发送SIGTERM
                    try:
                        os.killpg(os.getpgid(process.pid), signal.SIGTERM)
                        logger.info(f"[STOP] 发送SIGTERM到进程组 PID={process.pid}")
                    except (ProcessLookupError, PermissionError):
                        pass

                    # 直接终止进程
                    process.terminate()
                    logger.info(f"[STOP] 终止进程 PID={process.pid}")

                    # 等待一小段时间，如果还没结束就强制kill
                    try:
                        process.wait(timeout=2)
                    except subprocess.TimeoutExpired:
                        process.kill()
                        logger.info(f"[STOP] 强制结束进程 PID={process.pid}")

                    process_killed = True
                    del RUNNING_PROCESSES[request_id]
                except Exception as e:
                    logger.error(f"[STOP] 终止进程失败: {str(e)}")

        # 更新进度跟踪状态
        if request_id in PROGRESS_TRACKING:
            PROGRESS_TRACKING[request_id] = {
                "status": "stopped",
                "result": "用户手动停止了进程"
            }

        if process_killed:
            logger.info(f"[STOP] ✅ 进程已停止: request_id={request_id}")
            self.send_json_response(200, {
                "success": True,
                "message": "进程已停止",
                "request_id": request_id
            })
        else:
            logger.warning(f"[STOP] 进程未找到或已结束: request_id={request_id}")
            self.send_json_response(200, {
                "success": True,
                "message": "进程未找到或已结束",
                "request_id": request_id
            })

    def _handle_create_session(self, data):
        """创建会话"""
        session_id = str(uuid.uuid4())
        SESSIONS[session_id] = Session(session_id)

        logger.info(f"创建会话: {session_id}")
        self.send_json_response(200, {
            "session_id": session_id,
            "created_at": datetime.now().isoformat()
        })

    def _handle_send_message(self, data):
        """发送消息"""
        session_id = data.get('session_id')
        message = data.get('message', '')

        if not session_id or session_id not in SESSIONS:
            self.send_json_response(400, {"error": "Invalid session"})
            return

        session = SESSIONS[session_id]
        session.add_message('user', message)

        # 在线程池中处理
        request_id = str(uuid.uuid4())
        PROGRESS_TRACKING[request_id] = {"status": "processing"}

        def process():
            try:
                # 设置线程本地的request_id，供call_claude等函数使用
                CURRENT_REQUEST_CONTEXT.request_id = request_id

                logger.info(f"[PROCESS] 开始处理消息: {message}...")
                # 使用语义调度系统处理
                if semantic_scheduler is not None:
                    logger.info(f"[PROCESS] 使用语义调度系统处理")
                    result = semantic_scheduler.process(message, {"kb_path": KB_PATHS})
                    response = result["response"]
                    intent_info = result["intent"]
                    logger.info(f"[PROCESS] 语义调度完成，响应长度: {len(response)},intent:{str(result)}")
                else:
                    logger.info(f"[PROCESS] 使用直接调用模型处理")
                    # 回退到直接调用模型
                    response = call_model(message, model="nvidia/minimaxai/minimax-m2.1")
                    intent_info = {"type": "direct", "target": "claude"}
                    logger.info(f"[PROCESS] 直接调用完成，响应长度: {len(response)}")

                session.add_message('assistant', response)
                PROGRESS_TRACKING[request_id] = {
                    "status": "completed",
                    "result": response,
                    "intent": intent_info
                }
                logger.info(f"[PROCESS] 处理完成，request_id: {request_id}")
            except Exception as e:
                logger.error(f"[PROCESS] 处理异常: {e}", exc_info=True)
                PROGRESS_TRACKING[request_id] = {
                    "status": "error",
                    "error": str(e)
                }

        EXECUTOR.submit(process)

        self.send_json_response(202, {
            "request_id": request_id,
            "status": "processing"
        })

    def _handle_get_result(self):
        """获取异步处理结果"""
        parsed = urlparse(self.path)
        request_id = parsed.path.split("/")[-1]

        if request_id in PROGRESS_TRACKING:
            tracking = PROGRESS_TRACKING[request_id]
            # print(f"==tracking:{tracking}")

            if tracking != None and tracking["status"] == "completed":
                # 返回最终结果（包含HTTP状态码和数据）
                response_data = tracking.get("result", {})
                status_code = response_data.get("status_code", 200)
                response = response_data.get("data", {})

                # 发送最终结果，HTTP状态码仍为200（表示查询成功）
                self.send_json_response(200, {
                    "status": "completed",
                    "request_id": request_id,
                    "result": response,
                    "original_status_code": status_code,
                    "completed_at": tracking.get("end_time")
                })
            elif tracking != None and tracking["status"] == "running":
                # 仍在运行中
                self.send_json_response(202, {
                    "status": "running",
                    "request_id": request_id,
                    "message": "请求仍在处理中",
                    "started_at": tracking.get("start_time")
                })
            else:  # failed
                # 处理失败
                self.send_json_response(200, {
                    "status": "failed",
                    "request_id": request_id,
                    "error": tracking.get("error"),
                    "failed_at": tracking.get("end_time")
                })
        else:
            self.send_json_response(404, {
                "status": "not_found",
                "error": "请求未找到或已过期"
            })
#   扩展了 _handle_download_file 接口，支持多种下载方式：                                                                                                                          
#   ┌─────────────────────────────────┬──────────────────────┐                                                                                                                     
#   │            请求格式             │         说明         │                                                                                                                     
#   ├─────────────────────────────────┼──────────────────────┤                                                                                                                     
#   │ /api/download/KB/文件名         │ 从指定KB目录下载     │                                                                                                                     
#   ├─────────────────────────────────┼──────────────────────┤                                                                                                                     
#   │ /api/download/文件名            │ 自动在所有KB目录搜索 │                                                                                                                     
#   ├─────────────────────────────────┼──────────────────────┤                                                                                                                     
#   │ /api/download/opencode_p/文件名 │ 从下载目录下载       │                                                                                                                     
#   └─────────────────────────────────┴──────────────────────┘ 
#    1. 用户在前端输入命令                                                                                                                                                         
#       ↓                                                                                                                                                                          
#    2. 后端处理请求（scheduler.py）                                                                                                                                               
#       ↓                                                                                                                                                                          
#    3. 生成结果文件 → 保存到 /home/will/Downloads/opencode_p/tmp/                                                                                                                 
#       ↓                                                                                                                                                                          
#    4. 返回下载链接给前端                                                                                                                                                         
#       ↓                                                                                                                                                                          
#    5. 用户点击链接                                                                                                                                                               
#       ↓                                                                                                                                                                          
#    6. 后端搜索文件（server.py）                                                                                                                                                  
#       ├→ 优先搜索 /home/will/Downloads/opencode_p/tmp/                                                                                                                           
#       ├→ 再搜索知识库目录                                                                                                                                                        
#       └→ 最后搜索其他目录                                                                                                                                                        
#       ↓                                                                                                                                                                          
#    7. 找到文件并返回下载    
    def _handle_download_file(self, path):
        """处理文件下载请求 - /api/download/<kb>/<filename> 或 /api/download/<filename>"""
        import os
        from urllib.parse import unquote, quote

        # 提取文件路径信息
        path_parts = path.replace("/api/download/", "").strip().split("/")

        file_path = None
        search_locations = []  # 用于日志记录搜索位置

        # 情况1：知识库文件 /api/download/KB/filename
        if len(path_parts) == 2:
            kb_name = path_parts[0].upper()
            filename = unquote(path_parts[1])

            if kb_name in KB_PATHS:
                file_path = os.path.join(KB_PATHS[kb_name], filename)
                search_locations = [file_path]

        # 情况2：直接指定文件名 /api/download/filename
        elif len(path_parts) == 1:
            filename = unquote(path_parts[0])

            # 安全检查 - 防止路径遍历攻击
            if ".." in filename or filename.startswith("/"):
                self.send_json_response(400, {"error": "Invalid filename"})
                return

            # 定义搜索顺序
            search_dirs = [
                *KB_PATHS.values(),  # 所有知识库目录
                "/home/will/Downloads/opencode_p/tmp/",  # 新生成文件目录
                "/home/will/Downloads/opencode_p/",  # 下载目录
                WORK_DIR,  # 工作目录
            ]

            search_locations = []
            for kb_path in search_dirs:
                candidate = os.path.join(kb_path, filename)
                search_locations.append(candidate)
                if os.path.exists(candidate):
                    file_path = candidate
                    break

        else:
            self.send_json_response(400, {"error": "Invalid download path"})
            return

        # 安全检查 - 防止路径遍历攻击
        if ".." in str(file_path) or not file_path or not os.path.isabs(file_path):
            self.send_json_response(400, {"error": "Invalid filename"})
            return

        # 检查文件是否存在
        if not os.path.exists(file_path) or not os.path.isfile(file_path):
            logger.warning(f"[Download] 文件不存在: {file_path}")
            logger.warning(f"[Download] 搜索位置: {search_locations}")
            self.send_json_response(404, {"error": f"File not found: {filename}"})
            return

        try:
            # 读取文件内容
            with open(file_path, 'rb') as f:
                content = f.read()

            # 获取文件名
            basename = os.path.basename(file_path)

            # 设置响应头 - 根据RFC 5987正确处理中文文件名
            self.send_response(200)
            self.send_header('Content-Type', 'application/octet-stream')

            # 检查文件名是否包含非ASCII字符
            try:
                basename.encode('ascii')
                # 纯ASCII文件名，使用标准格式
                cd_header = f'attachment; filename="{basename}"'
            except UnicodeEncodeError:
                # 包含非ASCII字符（如中文），使用RFC 5987格式
                # 同时提供filename参数（兼容旧浏览器）和filename*参数（现代浏览器）
                encoded_filename = quote(basename, encoding='utf-8')
                cd_header = f'attachment; filename="{basename.encode("ascii", "replace").decode("ascii")}"; filename*=UTF-8\'\'{encoded_filename}'

            self.send_header('Content-Disposition', cd_header)
            self.send_header('Content-Length', len(content))
            self.end_headers()

            # 发送文件内容
            self.wfile.write(content)

            logger.info(f"[Download] ✅ 文件已下载: {file_path} ({len(content)} bytes)")

        except Exception as e:
            logger.error(f"[Download] 文件下载失败: {e}")
            self.send_json_response(500, {"error": f"Download failed: {e}"})

    def _handle_list_kb(self, data):
        """列出知识库文件"""
        kb_name = data.get('kb', 'KB')
        output = list_kb_files(kb_name)
        self.send_json_response(200, {"output": output})

    def _handle_search_kb(self, data):
        """搜索知识库"""
        query = data.get('query', '')
        kb_name = data.get('kb', 'KB')

        if not query:
            self.send_json_response(400, {"error": "Query required"})
            return

        output = search_kb(query, kb_name)
        self.send_json_response(200, {"output": output})

    def _handle_upload_url(self, data):
        """从URL提取内容保存为DOCX并索引"""
        url = data.get('url', '')
        kb_name = data.get('kb', 'KB')

        if not url:
            self.send_json_response(400, {"success": False, "error": "URL required"})
            return

        if kb_name not in KB_PATHS:
            self.send_json_response(400, {"success": False, "error": f"知识库'{kb_name}'不存在"})
            return

        try:
            result = extract_url_to_docx(url, kb_name)
            self.send_json_response(200, result)
        except Exception as e:
            logger.error(f"URL提取错误: {e}")
            self.send_json_response(500, {"success": False, "error": str(e)})

    def _handle_upload_file(self):
        """处理文档上传并索引"""
        try:
            content_type = self.headers.get('Content-Type', '')
            if 'multipart/form-data' not in content_type:
                self.send_json_response(400, {"success": False, "error": "需要 multipart/form-data"})
                return

            # 解析 multipart 数据
            form = cgi.FieldStorage(
                fp=self.rfile,
                headers=self.headers,
                environ={
                    'REQUEST_METHOD': 'POST',
                    'CONTENT_TYPE': content_type,
                }
            )

            file_item = form['file']
            kb_name = form.getvalue('kb', 'KB')

            if not file_item.filename:
                self.send_json_response(400, {"success": False, "error": "未选择文件"})
                return

            if kb_name not in KB_PATHS:
                self.send_json_response(400, {"success": False, "error": f"知识库'{kb_name}'不存在"})
                return

            file_data = file_item.file.read()
            result = save_uploaded_file(file_data, file_item.filename, kb_name)
            self.send_json_response(200, result)
        except Exception as e:
            logger.error(f"文件上传错误: {e}")
            self.send_json_response(500, {"success": False, "error": str(e)})

    def _handle_read_file(self, data):
        """读取文件"""
        filepath = data.get('path', '')

        if not filepath:
            self.send_json_response(400, {"error": "Path required"})
            return

        content = read_file(filepath)
        self.send_json_response(200, {"content": content})

    def _handle_send_content_email(self, data):
        """发送文本内容到邮箱"""
        to_email = data.get('to_email', '')
        subject = data.get('subject', '来自Web服务器的邮件')
        content = data.get('content', '')
        convert_to_docx = data.get('convert_to_docx', False)

        if not to_email or not content:
            self.send_json_response(400, {"success": False, "error": "邮箱和内容不能为空"})
            return

        try:
            attachment_path = None
            if convert_to_docx:
                # 转换为 DOCX 文件
                title = subject.replace('/', '-').replace('\\', '-')[:30]
                attachment_path = content_to_docx(content, title)
                if not attachment_path:
                    self.send_json_response(500, {"success": False, "error": "DOCX 生成失败"})
                    return

            # 发送邮件
            success = send_email(to_email, subject, content, attachment_path)
            if success:
                self.send_json_response(200, {"success": True, "message": f"邮件已发送到 {to_email}"})
            else:
                self.send_json_response(500, {"success": False, "error": "邮件发送失败，请检查邮箱配置"})
        except Exception as e:
            logger.error(f"[Email] 发送失败: {e}")
            self.send_json_response(500, {"success": False, "error": str(e)})

    def _handle_send_file_email(self, data):
        """发送文件到邮箱"""
        to_email = data.get('to_email', '')
        subject = data.get('subject', '文件分享')
        filepath = data.get('filepath', '')

        if not to_email or not filepath:
            self.send_json_response(400, {"success": False, "error": "邮箱和文件路径不能为空"})
            return

        if not os.path.exists(filepath):
            self.send_json_response(400, {"success": False, "error": f"文件不存在: {filepath}"})
            return

        try:
            body = f"请查看附件中的文件: {os.path.basename(filepath)}"
            success = send_email(to_email, subject, body, filepath)
            if success:
                self.send_json_response(200, {"success": True, "message": f"文件已发送到 {to_email}"})
            else:
                self.send_json_response(500, {"success": False, "error": "邮件发送失败"})
        except Exception as e:
            logger.error(f"[Email] 发送文件失败: {e}")
            self.send_json_response(500, {"success": False, "error": str(e)})

    def _handle_send_kb_email(self, data):
        """发送知识库文件到邮箱"""
        to_email = data.get('to_email', '')
        kb_name = data.get('kb', 'KB')
        filename = data.get('filename', '')

        if not to_email or not filename:
            self.send_json_response(400, {"success": False, "error": "邮箱和文件名不能为空"})
            return

        if kb_name not in KB_PATHS:
            self.send_json_response(400, {"success": False, "error": f"知识库 '{kb_name}' 不存在"})
            return

        try:
            filepath = Path(KB_PATHS[kb_name]) / filename
            if not filepath.exists():
                self.send_json_response(400, {"success": False, "error": f"文件不存在: {filename}"})
                return

            subject = f"知识库文件: {filename}"
            body = f"来自 {kb_name} 知识库的文件"
            success = send_email(to_email, subject, body, str(filepath))

            if success:
                self.send_json_response(200, {"success": True, "message": f"文件已从 {kb_name} 发送到 {to_email}"})
            else:
                self.send_json_response(500, {"success": False, "error": "邮件发送失败"})
        except Exception as e:
            logger.error(f"[KB Email] 发送失败: {e}")
            self.send_json_response(500, {"success": False, "error": str(e)})

    def _handle_send_article_email(self, data):
        """发送文章到邮箱（支持 Markdown 格式，自动转换为 DOCX）

        参数:
            to_email: 收件人邮箱（必填）
            subject: 邮件主题（必填）
            article_content: 文章内容（Markdown 格式，必填）
            title: 文档标题，不指定时使用 subject（可选）
            email_body: 邮件正文，不指定时使用默认文本（可选）
        """
        to_email = data.get('to_email', '')
        subject = data.get('subject', '')
        article_content = data.get('article_content', '')
        doc_title = data.get('title', subject or '文章')
        email_body = data.get('email_body', f'尊敬的用户，\n\n您好！附件中是文章 "{subject}" 的 DOCX 版本。\n\n祝好！\nWeb 服务器')

        # 参数验证
        if not to_email:
            self.send_json_response(400, {"success": False, "error": "收件人邮箱不能为空"})
            return

        if not subject:
            self.send_json_response(400, {"success": False, "error": "邮件主题不能为空"})
            return

        if not article_content:
            self.send_json_response(400, {"success": False, "error": "文章内容不能为空"})
            return

        try:
            # 第1步：将 Markdown 内容转换为 DOCX
            logger.info(f"[Article Email] 开始生成 DOCX 文档: {doc_title}")
            docx_path = content_to_docx(article_content, doc_title)

            if not docx_path:
                self.send_json_response(500, {"success": False, "error": "DOCX 文档生成失败"})
                return

            logger.info(f"[Article Email] DOCX 文档生成成功: {docx_path}")

            # 第2步：发送邮件
            logger.info(f"[Article Email] 准备发送邮件到 {to_email}")
            success = send_email(to_email, subject, email_body, docx_path)

            if success:
                logger.info(f"[Article Email] ✅ 邮件发送成功: {to_email}")
                self.send_json_response(200, {
                    "success": True,
                    "message": f"文章已成功发送到 {to_email}",
                    "docx_path": docx_path,
                    "file_size": os.path.getsize(docx_path) if os.path.exists(docx_path) else 0
                })
            else:
                logger.error(f"[Article Email] ❌ 邮件发送失败")
                self.send_json_response(500, {"success": False, "error": "邮件发送失败，请检查邮箱配置"})

        except Exception as e:
            logger.error(f"[Article Email] ❌ 异常: {e}")
            self.send_json_response(500, {"success": False, "error": str(e)})

    def _analyze_intent(self, prompt: str) -> dict:
        """分析用户意图，返回意图类型和参数"""
        prompt_lower = prompt.lower().strip()
        original_prompt = prompt.strip()

        # 1. 意图：知识库文件列表（最高优先级）
        # 检查列表关键词 + 知识库/文件上下文
        list_keywords = ['列出', '列表', '显示', '展示', 'list', 'show', 'ls']
        kb_context_words = ['知识库', 'kb', '文件', '文档', 'documents', 'files']

        # 检查是否包含列表关键词
        has_list_keyword = any(keyword in prompt_lower for keyword in list_keywords)
        has_kb_context = any(context in prompt_lower for context in kb_context_words)

        if has_list_keyword and has_kb_context:
            # 确定知识库名称
            kb_name = "KB"  # 默认
            if 'kbgx' in prompt_lower or '国学' in prompt_lower:
                kb_name = "KBGX"
            elif 'kbw' in prompt_lower or '微信' in prompt_lower:
                kb_name = "KBW"
            elif 'kb' in prompt_lower:
                # 检查是否指定了特定的KB
                if 'kb ' in prompt_lower:
                    # 简单提取：查找"KB"后的字符
                    kb_match = re.search(r'kb\s*(\w+)', prompt_lower)
                    if kb_match:
                        kb_suffix = kb_match.group(1)
                        if 'gx' in kb_suffix:
                            kb_name = "KBGX"
                        elif 'w' in kb_suffix:
                            kb_name = "KBW"

            return {
                "intent": "kb_list",
                "kb": kb_name
            }

        # # 2. 意图：Web搜索
        # web_search_keywords = ['从网络搜索', '网络搜索', 'web搜索', 'web search', '从网络查找', '互联网搜索']
        # for keyword in web_search_keywords:
        #     if keyword in prompt_lower:
        #         # 提取查询内容（保留原始大小写）
        #         query = original_prompt
        #         # 找到关键词位置并移除
        #         idx = prompt_lower.find(keyword)
        #         if idx != -1:
        #             query = original_prompt[idx + len(keyword):].strip()
        #         if not query:
        #             query = original_prompt
        #         return {
        #             "intent": "web_search",
        #             "query": query
        #         }

        # 3. 意图：知识库搜索
        # 关键：首先检查"搜索"关键字，排除已被识别为列表请求的情况
        kb_search_keywords = ['搜索', '查找', '查询', 'search', 'find', '检索', '总结', '汇总']

        # 搜索请求优先于列表请求处理
        if any(kw in prompt_lower for kw in kb_search_keywords):
            # 检查是否包含网络搜索排除词（如果是，已经是web_search）
            web_exclude = False
            # for web_kw in web_search_keywords:
            #     if web_kw in prompt_lower:
            #         web_exclude = True
            #         break

            if not web_exclude:
                # 找到搜索关键词
                for keyword in kb_search_keywords:
                    if keyword in prompt_lower:
                        # 提取查询内容
                        query = original_prompt
                        idx = prompt_lower.find(keyword)
                        if idx != -1:
                            query = original_prompt[idx + len(keyword):].strip()
                        if not query:
                            query = original_prompt

                        # 确定知识库
                        kb_name = "KB"
                        if 'kbgx' in prompt_lower or '国学' in prompt_lower:
                            kb_name = "KBGX"
                        elif 'kbw' in prompt_lower or '微信' in prompt_lower:
                            kb_name = "KBW"

                        return {
                            "intent": "kb_search",
                            "query": query,
                            "kb": kb_name
                        }

        # 4. 意图：读取文件
        file_keywords = ['读取文件', '打开文件', '查看文件', 'read file', 'open file', 'file read']
        for keyword in file_keywords:
            if keyword in prompt_lower:
                return {
                    "intent": "needs_analysis",
                    "reason": "可能需要文件操作，需要进一步解析文件路径"
                }

        # 默认：需要LLM直接回答
        return {
            "intent": "direct_answer",
            "prompt": original_prompt
        }
    #后端接收输入，模型，知识库
    def _handle_claude_call(self, data):
        """智能调用Claude - 使用语义调度系统路由到不同的 Skill/Agent"""
        prompt = data.get('prompt', '')
        model = data.get('model', 'haiku')
        frontend_kb = data.get('kb', None)
        history = data.get('history', [])  # 获取对话历史
        use_semantic_scheduler = True #前端没有发送 data.get('use_semantic_scheduler', True)  # 开关控制

        if not prompt:
            self.send_json_response(400, {"error": "Prompt required"})
            return

        # 记录对话历史信息
        if history:
            logger.info(f"[Claude] 收到 {len(history)} 条对话历史")

        # 如果启用语义调度系统且已初始化，使用它来处理
        if use_semantic_scheduler and semantic_scheduler is not None:
            logger.info(f"[Claude] 使用语义调度系统处理: {prompt}...")
            try:  #result={'response':返回的知识库列表文件，'intent': {'type': 'skill', 'target': 'knowledge-base', 'operation': 'list_docs', 'reasoning': '检测到知识库文件列表请求'}, 'metadata': {'processing_time': 5.119239, 'confidence': 1.0, 'available_skills': 21, 'available_agents': 5}}
                # 获取服务器的base_url，用于生成完整的下载链接
                host = self.headers.get('Host', 'localhost:8085')
                base_url = f"http://{host}"
                # 传递对话历史和base_url给语义调度系统
                result = semantic_scheduler.process(prompt, {"kb": frontend_kb, "model": model, "history": history, "base_url": base_url})
                self.send_json_response(200, {
                    "response": result["response"],
                    "intent": result["intent"],
                    "metadata": result["metadata"],
                    "scheduler_used": True
                })
                return
            except Exception as e:
                logger.warning(f"[Claude] 语义调度失败: {e}，使用默认处理")

        # 默认：使用原有的关键词匹配意图分析
        intent = self._analyze_intent(prompt)
        logger.info(f"[Claude] 使用关键词匹配: {intent}")

        # 根据意图路由
        if intent["intent"] == "kb_search":
            query = intent.get("query", prompt)
            kb = frontend_kb if frontend_kb else intent.get("kb", "KB")
            output = search_kb(query, kb)
            self.send_json_response(200, {"response": output, "intent": "kb_search", "scheduler_used": False})

        elif intent["intent"] == "kb_list":
            kb = frontend_kb if frontend_kb else intent.get("kb", "KB")
            output = list_kb_files(kb)
            self.send_json_response(200, {"response": output, "intent": "kb_list", "scheduler_used": False})
        #这是错误的，call claude todo
        # elif intent["intent"] == "web_search":
        #     query = intent.get("query", prompt)
        #     response = call_model(
        #         f"请基于您的知识库搜索并提供关于\"{query}\"的最新信息总结。\n\n请包括：\n1. 主要信息和进展\n2. 关键数据和日期\n3. 相关的公司或机构\n4. 最近的动态或计划\n\n请以清晰、结构化的格式提供信息。",
        #         model=model
        #     )
        #     self.send_json_response(200, {"response": response, "intent": "web_search", "scheduler_used": False})

        elif intent["intent"] == "direct_answer":
            response = call_model(prompt, model=model)
            self.send_json_response(200, {"response": response, "intent": "direct_answer", "scheduler_used": False})

        else:
            response = call_model(prompt, model=model)
            self.send_json_response(200, {"response": response, "intent": "unknown", "scheduler_used": False})

    def _handle_web_search(self, data):
        """处理Web搜索请求"""
        query = data.get('query', '')
        model = 'haiku' #data.get('model', 'glm4.7')

        if not query:
            logger.warning(f"[WEB] 缺少查询参数")
            self.send_json_response(400, {"error": "Query required"})
            return

        logger.info(f"[WEB] 开始搜索: '{query}' (使用模型: {model})")

        # 构建搜索提示 - 更明确、更强制地要求执行搜索
        search_prompt = f"""{query}

立即执行以下操作（不需要确认）：
1. 使用WebSearch工具搜索相关信息
2. 整理搜索结果为清晰的回答
3. 包含信息来源链接

要求：
- 直接执行搜索，不需要用户确认
- 提供最新、最相关的信息
- 包含出处和时间信息
- 中文回答"""

        try:
            logger.info(f"[WEB] 调用Claude API开始处理...")
            # Web搜索需要更长的超时时间（包括Claude API调用 + WebSearch工具执行）
            response = call_model(search_prompt, model=model, timeout=600)
            logger.info(f"[WEB] ✅ 搜索完成，结果长度: {len(response)}")
            self.send_json_response(200, {"results": response, "query": query})
        except Exception as e:
            logger.error(f"[WEB] ❌ 搜索失败: {e}", exc_info=True)
            self.send_json_response(500, {"error": str(e)})

    def _handle_web_search_dk(self, data):
        """使用 dk-search 服务进行网络搜索"""
        import httpx
        from datetime import datetime

        query = data.get('query', '')
        provider = data.get('provider', 'ddg')  # ddg 或 google
        max_results = data.get('max_results', 5)

        if not query:
            logger.warning(f"[WEB-DK] 缺少查询参数")
            self.send_json_response(400, {"error": "Query required"})
            return

        # 检查查询中是否包含日期/时间相关词汇
        date_keywords = ['今天', '今日', '昨天', '本周', '本月', '最近', '最新',
                        '2024', '2025', '2026', '2027', '年', '月', '日',
                        'today', 'yesterday', 'this week', 'this month', 'recent', 'latest']
        has_date = any(keyword in query.lower() for keyword in date_keywords)

        # 如果没有日期信息，添加当前日期
        if not has_date:
            current_date = datetime.now().strftime("%Y年%m月")
            query = f"{query} {current_date}"
            logger.info(f"[WEB-DK] 自动添加日期，查询变更为: '{query}'")

        logger.info(f"[WEB-DK] 开始搜索: '{query}' (provider: {provider}, max_results: {max_results})")

        dk_search_url = "http://localhost:8001/fetch"

        try:
            with httpx.Client(timeout=60) as client:
                response = client.post(
                    dk_search_url,
                    json={
                        "query": query,
                        "provider": provider,
                        "max_results": max_results,
                        "timeout": 15
                    }
                )

                if response.status_code != 200:
                    logger.error(f"[WEB-DK] dk-search 返回错误: {response.status_code}")
                    self.send_json_response(502, {
                        "error": f"dk-search service error: {response.status_code}",
                        "details": response.text
                    })
                    return

                result = response.json()

                # 格式化结果
                formatted_results = {
                    "query": result.get("query"),
                    "provider": result.get("provider"),
                    "results_count": len(result.get("results", [])),
                    "results": [
                        {
                            "title": r.get("title"),
                            "url": r.get("url"),
                            "snippet": r.get("snippet"),
                            "content": r.get("content", "")[:2000] + "..." if len(r.get("content", "")) > 2000 else r.get("content", ""),
                            "content_chars": r.get("content_chars", 0)
                        }
                        for r in result.get("results", [])
                    ]
                }

                logger.info(f"[WEB-DK] ✅ 搜索完成，获取 {formatted_results['results_count']} 条结果")
                self.send_json_response(200, formatted_results)

        except httpx.ConnectError:
            logger.error(f"[WEB-DK] ❌ 无法连接到 dk-search 服务 (localhost:8001)")
            self.send_json_response(503, {
                "error": "dk-search service unavailable",
                "hint": "请确保 dk-search 服务正在运行 (端口 8001)"
            })
        except Exception as e:
            logger.error(f"[WEB-DK] ❌ 搜索失败: {e}", exc_info=True)
            self.send_json_response(500, {"error": str(e)})

    def _handle_list_skills(self, data):
        """列出所有可用的 Skills"""
        if semantic_scheduler is None:
            self.send_json_response(500, {"error": "Semantic scheduler not initialized"})
            return

        skills = []
        for name, skill in semantic_scheduler.discovery.skills.items():
            skills.append({
                "name": skill.name,
                "description": skill.description[:200] + "..." if len(skill.description) > 200 else skill.description,
                "operations": list(skill.operations.keys()),
                "keywords": skill.keywords,
                "path": str(skill.path)
            })

        self.send_json_response(200, {
            "count": len(skills),
            "skills": skills
        })

    def _handle_list_agents(self, data):
        """列出所有可用的 Agents"""
        if semantic_scheduler is None:
            self.send_json_response(500, {"error": "Semantic scheduler not initialized"})
            return

        agents = []
        for name, agent in semantic_scheduler.discovery.agents.items():
            agents.append({
                "name": agent.name,
                "description": agent.description,
                "subagent_type": agent.subagent_type,
                "use_cases": agent.use_cases
            })

        self.send_json_response(200, {
            "count": len(agents),
            "agents": agents
        })

    def _handle_skills_summary(self, data):
        """获取 Skills 和 Agents 的摘要信息"""
        if semantic_scheduler is None:
            self.send_json_response(500, {"error": "Semantic scheduler not initialized"})
            return

        summary = semantic_scheduler.discovery.get_skill_summary()
        self.send_json_response(200, summary)

    def send_json_response(self, status_code, data):
        """发送JSON响应"""
        try:
            self.send_response(status_code)
            self.send_header('Content-Type', 'application/json')
            self.send_header('Access-Control-Allow-Origin', '*')
            self.end_headers()
            self.wfile.write(json.dumps(data).encode())
        except BrokenPipeError:
            logger.warning(f"客户端断开连接，无法发送响应")
        except Exception as e:
            logger.error(f"发送响应时出错: {str(e)}")

    def send_html_response(self, html):
        """发送HTML响应"""
        self.send_response(200)
        self.send_header('Content-Type', 'text/html; charset=utf-8')
        self.end_headers()
        self.wfile.write(html.encode('utf-8'))

#  1. removeThinkTags() - 用字符串查找替换代替正则表达式                                                                                                                          
                                                                                                                                                                                 
#   // 移除 <think>...</think> 标签                                                                                                                                                
#   var thinkStart = text.indexOf('<think>');                                                                                                                                      
#   while (thinkStart >= 0) {                                                                                                                                                      
#       var thinkEnd = result.indexOf('</think>', thinkStart);                                                                                                                     
#       if (thinkEnd >= 0) {                                                                                                                                                       
#           result = result.substring(0, thinkStart) + result.substring(thinkEnd + 8);                                                                                             
#       }                                                                                                                                                                          
#       thinkStart = result.indexOf('<think>');                                                                                                                                    
#   }                                                                                                                                                                              
                                                                                                                                                                                 
#   2. linkifyUrls() - 用字符串查找替换URL                                                                                                                                         
                                                                                                                                                                                 
#   // 查找 http:// 或 https://                                                                                                                                                    
#   var httpStart = result.indexOf('http://');                                                                                                                                     
#   if (httpStart < 0) httpStart = result.indexOf('https://');                                                                                                                     
#   // 逐个转换成链接                                                                                                                                                                 #    
                  
                  
    def get_frontend_html(self):
        """返回前端HTML - 改进版布局"""
        return """
<!DOCTYPE html>
<html lang="zh-CN">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Claude Web Server LLM</title>
    <style>
        * { margin: 0; padding: 0; box-sizing: border-box; }
        html, body {
            width: 100%;
            height: 100%;
            font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Menlo, monospace;
        }
        body {
            background: #000000;
            color: #e0e0e0;
            display: flex;
            flex-direction: column;
        }

        .header {
            background: #1a1a1a;
            border-bottom: 1px solid #333;
            padding: 12px 20px;
            display: flex;
            justify-content: space-between;
            align-items: center;
            min-height: 50px;
        }

        .header h1 {
            font-size: 18px;
            color: #4a9eff;
            font-weight: bold;
        }

        .status {
            font-size: 12px;
            color: #888;
        }

        .main-container {
            display: flex;
            flex: 1;
            width: 100%;
            overflow: hidden;
        }

        .messages-container {
            flex: 1;
            display: flex;
            flex-direction: column;
            padding: 20px;
            overflow: hidden;
        }

        .messages {
            flex: 1;
            overflow-y: auto;
            display: flex;
            flex-direction: column;
            gap: 12px;
            margin-bottom: 20px;
            padding-right: 10px;
        }

        .messages::-webkit-scrollbar {
            width: 8px;
        }

        .messages::-webkit-scrollbar-track {
            background: #111;
        }

        .messages::-webkit-scrollbar-thumb {
            background: #333;
            border-radius: 4px;
        }

        .message {
            padding: 10px 14px;
            border-radius: 6px;
            word-wrap: break-word;
            white-space: pre-wrap;
            line-height: 1.5;
            max-width: 100%;
        }

        .message.user {
            align-self: flex-end;
            background: #0066cc;
            color: white;
            max-width: 85%;
            border-radius: 12px 4px 12px 12px;
        }

        .message.assistant {
            align-self: flex-start;
            background: #222;
            color: #e0e0e0;
            max-width: 100%;
            border-left: 3px solid #4a9eff;
        }

        .message.system {
            align-self: center;
            background: #333;
            color: #aaa;
            font-size: 12px;
            padding: 6px 10px;
            border-radius: 4px;
        }

        .message.loading {
            align-self: flex-start;
            background: #222;
            color: #888;
            max-width: 100px;
            display: flex;
            align-items: center;
            gap: 8px;
            font-size: 12px;
            border-left: 3px solid #4a9eff;
        }

        .spinner {
            width: 16px;
            height: 16px;
            border: 2px solid #333;
            border-top: 2px solid #4a9eff;
            border-radius: 50%;
            animation: spin 1s linear infinite;
        }

        @keyframes spin {
            0% { transform: rotate(0deg); }
            100% { transform: rotate(360deg); }
        }

        .input-section {
            display: flex;
            flex-direction: column;
            gap: 12px;
            border-top: 1px solid #333;
            padding-top: 12px;
        }

        .input-controls {
            display: flex;
            gap: 10px;
        }

        textarea {
            flex: 1;
            background: #1a1a1a;
            border: 1px solid #333;
            border-radius: 6px;
            padding: 10px 12px;
            color: #e0e0e0;
            font-size: 14px;
            font-family: 'Menlo', monospace;
            resize: none;
            min-height: 60px;
            max-height: 120px;
        }

        textarea::placeholder {
            color: #666;
        }

        textarea:focus {
            outline: none;
            border-color: #4a9eff;
            box-shadow: 0 0 0 2px rgba(74, 158, 255, 0.1);
        }

        .send-btn {
            background: #0066cc;
            color: white;
            border: none;
            border-radius: 6px;
            padding: 10px 20px;
            cursor: pointer;
            font-weight: bold;
            min-width: 70px;
            transition: background 0.2s;
        }

        .send-btn:hover {
            background: #0052a3;
        }

        .quick-actions {
            display: flex;
            flex-wrap: wrap;
            gap: 8px;
        }

        .message-toolbar {
            display: flex;
            gap: 8px;
            margin-top: 8px;
            opacity: 0;
            transition: opacity 0.2s;
        }

        .message:hover .message-toolbar {
            opacity: 1;
        }

        .download-btn {
            background: #2a5a8a;
            color: #ddd;
            border: 1px solid #3a6a9a;
            border-radius: 4px;
            padding: 4px 10px;
            cursor: pointer;
            font-size: 11px;
            transition: all 0.2s;
        }

        .download-btn:hover {
            background: #3a6a9a;
            border-color: #4a7aaa;
            color: #fff;
        }

        .quick-btn {
            background: #222;
            color: #ddd;
            border: 1px solid #333;
            border-radius: 4px;
            padding: 6px 12px;
            cursor: pointer;
            font-size: 12px;
            transition: all 0.2s;
        }

        .quick-btn:hover {
            background: #333;
            border-color: #4a9eff;
            color: #4a9eff;
        }

        .controls-row {
            display: flex;
            gap: 12px;
            flex-wrap: wrap;
        }

        .control-group {
            display: flex;
            gap: 8px;
            align-items: center;
            font-size: 12px;
        }

        .control-group label {
            color: #999;
        }

        select {
            background: #1a1a1a;
            color: #e0e0e0;
            border: 1px solid #333;
            border-radius: 4px;
            padding: 6px 8px;
            font-size: 12px;
            cursor: pointer;
        }

        select:focus {
            outline: none;
            border-color: #4a9eff;
        }
    </style>
</head>
<body>
    <div class="header">
        <h1>Claude Web Server LLM</h1>
        <span class="status" id="status">准备就绪</span>
    </div>

    <div class="main-container">
        <div class="messages-container" style="width: 100%; padding-left: max(20px, calc((100% - 900px) / 2)); padding-right: max(20px, calc((100% - 900px) / 2));">
            <div class="messages" id="messages"></div>

            <div class="input-section">
                <div class="input-controls">
                    <textarea id="input" placeholder="输入您的问题... (Shift+Enter 快速发送)" rows="3"></textarea>
                    <button class="send-btn" onclick="sendMessage()">发送</button>
                </div>

                <div class="controls-row">
                    <div class="control-group">
                        <label>知识库:</label>
                        <select id="kb-select">
                            <option value="KB">通用库</option>
                            <option value="KBGX">国学库</option>
                            <option value="KBW">微信库</option>
                        </select>
                    </div>

                    <div class="control-group">
                        <label>模型:</label>
                        <select id="model-select">
                            <option value="minimax-m2.1">MiniMax M2.1</option>
                            <option value="deepseek-v3.2">Deepseek V3.2</option>
                            <option value="deepseek-v3.1-terminus">Deepseek V3.1</option>
                            <option value="haiku">Claude Haiku (快速)</option>
                            <option value="sonnet">Claude Sonnet (推荐)</option>
                            <option value="opus">Claude Opus (最强)</option>
                            <option value="glm4.7">GLM4.7</option>
                            <option value="deepseek-chat">Deepseek Chat</option>
                        </select>
                    </div>
                </div>

                <div class="quick-actions">
                    <button class="quick-btn" onclick="testConnection()">🔗 测试连接</button>
                    <button class="quick-btn" onclick="listKB()">📚 列出知识库</button>
                    <button class="quick-btn" onclick="searchKB()">🔍 搜索知识库</button>
                    <button class="quick-btn" onclick="uploadURL()">📤 上传URL</button>
                    <button class="quick-btn" onclick="uploadDoc()">📁 上传文档</button>
                    <button class="quick-btn" onclick="webSearch()">🌐 Web搜索</button>
                    <button class="quick-btn" onclick="clearMessages()">🗑️ 清空聊天</button>
                    <button class="quick-btn" onclick="stopProcess()" id="stop-btn" style="display: none; color: #ff6b6b;">⏹️ 停止进程</button>
                    <input type="file" id="file-upload" style="display: none;"
                           accept=".pdf,.docx,.xlsx,.pptx,.epub,.md,.txt,.json,.yaml,.yml,.csv,.png,.jpg,.jpeg,.mp3,.wav,.html"
                           onchange="handleFileUpload(event)">
                </div>
            </div>
        </div>
    </div>

    <script>
        let sessionId = null;
        // 对话历史：存储最近的对话（最多保留3轮，即6条消息）
        let conversationHistory = [];
        const MAX_HISTORY_ROUNDS = 3;  // 最多保留3轮对话

        // 任务控制变量
        let currentRequestId = null;  // 当前请求ID
        let shouldStopPolling = false;  // 是否停止轮询

        // 判断是否是新话题（需要丢弃历史上下文）
        function isNewTopic(message) {
            const msgLower = message.toLowerCase();
            // 新话题关键词：搜索、查找、列出、新问题等
            const newTopicKeywords = [
                '搜索', '查找', '检索', '列出', '显示文件', '文件列表',
                '新问题', '换个话题', '重新开始', '清空上下文',
                '网络搜索', 'web搜索', '今天', '最新', '天气', '新闻',
                '帮我', '请问', '什么是', '如何', '怎么', '为什么'
            ];

            // 如果消息以问号开头或包含新话题关键词，认为是新话题
            if (msgLower.startsWith('?') || msgLower.startsWith('？')) {
                return true;
            }

            // 检查是否包含新话题关键词
            for (const keyword of newTopicKeywords) {
                if (msgLower.includes(keyword)) {
                    return true;
                }
            }

            return false;
        }

        // 判断是否是引用上文的请求（需要保留历史上下文）
        function isContextualRequest(message) {
            const msgLower = message.toLowerCase();
            // 引用上文关键词
            const contextualKeywords = [
                '上面', '上述', '刚才', '前面', '这个', '那个',
                '转换', '导出', '下载', '保存', '生成',
                '继续', '接着', '然后', '另外', '还有',
                'word', 'docx', 'ppt', 'pptx', 'md', 'markdown',
                '第一', '第二', '第三', '其中', '哪个', '哪些'
            ];

            for (const keyword of contextualKeywords) {
                if (msgLower.includes(keyword)) {
                    return true;
                }
            }

            return false;
        }

        // 获取相关的对话历史（用于附加到请求中）
        function getRelevantHistory(currentMessage) {
            // 如果是新话题，不返回历史
            if (isNewTopic(currentMessage) && !isContextualRequest(currentMessage)) {
                console.log('[Context] 检测到新话题，清空历史上下文');
                return [];
            }

            // 如果是引用上文的请求，返回最近的历史
            if (isContextualRequest(currentMessage) || conversationHistory.length > 0) {
                console.log(`[Context] 附加 ${conversationHistory.length} 条历史消息`);
                return conversationHistory.slice();  // 返回副本
            }

            return [];
        }

        // 添加消息到历史记录
        function addToHistory(role, content) {
            // 不记录系统消息
            if (role === 'system') return;

            conversationHistory.push({ role, content });

            // 保留最近 MAX_HISTORY_ROUNDS 轮对话（每轮2条消息）
            const maxMessages = MAX_HISTORY_ROUNDS * 2;
            if (conversationHistory.length > maxMessages) {
                conversationHistory = conversationHistory.slice(-maxMessages);
            }

            console.log(`[History] 当前历史: ${conversationHistory.length} 条消息`);
        }

        // 清空对话历史
        function clearHistory() {
            conversationHistory = [];
            console.log('[History] 历史已清空');
        }

        async function initSession() {
            try {
                const response = await fetch('/api/session', { method: 'POST' });
                const data = await response.json();
                sessionId = data.session_id;
                addMessage('system', '✅ 会话已创建，可以开始对话');
            } catch (e) {
                addMessage('system', '❌ 连接失败: ' + e.message);
            }
        }

        async function testConnection() {
            try {
                const response = await fetch('/api/status');
                const data = await response.json();
                const msg = `✅ 服务器正常\\n• 状态: ${data.status}\\n• 活跃会话: ${data.sessions}`;
                addMessage('system', msg);
            } catch (e) {
                addMessage('system', '❌ 服务器无响应');
            }
        }

        async function listKB() {
            const kb = document.getElementById('kb-select').value;
            addMessage('system', `⏳ 正在列出${kb}文件...`);

            try {
                const response = await fetch('/api/kb/list', {
                    method: 'POST',
                    headers: { 'Content-Type': 'application/json' },
                    body: JSON.stringify({ kb })
                });
                const data = await response.json();
                if (data.output !== undefined && data.output !== null) {
                    addMessage('assistant', data.output);
                } else if (data.error) {
                    addMessage('system', '❌ 列出失败: ' + data.error);
                } else {
                    addMessage('system', '❌ 列出失败: 无效的响应格式');
                }
            } catch (e) {
                addMessage('system', '❌ 列出失败: ' + e.message);
            }
        }

        async function searchKB() {
            const query = prompt('输入搜索关键词:');
            if (!query) return;

            const kb = document.getElementById('kb-select').value;
            addMessage('user', `搜索: ${query}`);
            addMessage('system', `⏳ 正在搜索${kb}...`);

            try {
                const response = await fetch('/api/kb/search', {
                    method: 'POST',
                    headers: { 'Content-Type': 'application/json' },
                    body: JSON.stringify({ query, kb })
                });
                const data = await response.json();
                if (data.output !== undefined && data.output !== null) {
                    addMessage('assistant', data.output);
                } else if (data.error) {
                    addMessage('system', '❌ 搜索失败: ' + data.error);
                } else {
                    addMessage('system', '❌ 搜索失败: 无效的响应格式');
                }
            } catch (e) {
                addMessage('system', '❌ 搜索失败: ' + e.message);
            }
        }

        async function uploadURL() {
            const url = prompt('输入要提取内容的URL:');
            if (!url) return;

            const kb = document.getElementById('kb-select').value;
            addMessage('user', `上传URL: ${url}`);
            addMessage('system', `⏳ 正在提取URL内容并保存到${kb}...`);

            const loadingMsg = addLoadingMessage();

            try {
                const response = await fetch('/api/kb/upload-url', {
                    method: 'POST',
                    headers: { 'Content-Type': 'application/json' },
                    body: JSON.stringify({ url, kb })
                });
                const data = await response.json();
                removeMessage(loadingMsg);

                if (data.success) {
                    let msg = `✅ ${data.message}\\n📄 文件: ${data.filename}`;
                    msg += data.indexed ? '\\n📊 已自动索引' : '\\n⚠️ 索引未完成';
                    addMessage('assistant', msg);
                } else {
                    addMessage('system', '❌ 上传失败: ' + data.error);
                }
            } catch (e) {
                removeMessage(loadingMsg);
                addMessage('system', '❌ 上传失败: ' + e.message);
            }
        }

        function uploadDoc() {
            document.getElementById('file-upload').click();
        }

        async function handleFileUpload(event) {
            const file = event.target.files[0];
            if (!file) return;

            const kb = document.getElementById('kb-select').value;
            addMessage('user', `上传文档: ${file.name}`);
            addMessage('system', `⏳ 正在上传文档到${kb}并索引...`);

            const loadingMsg = addLoadingMessage();

            try {
                const formData = new FormData();
                formData.append('file', file);
                formData.append('kb', kb);

                const response = await fetch('/api/kb/upload-file', {
                    method: 'POST',
                    body: formData
                });
                const data = await response.json();
                removeMessage(loadingMsg);

                if (data.success) {
                    let msg = `✅ ${data.message}\\n📄 文件: ${data.filename}`;
                    msg += data.indexed ? '\\n📊 已自动索引' : '\\n⚠️ 索引未完成';
                    addMessage('assistant', msg);
                } else {
                    addMessage('system', '❌ 上传失败: ' + data.error);
                }
            } catch (e) {
                removeMessage(loadingMsg);
                addMessage('system', '❌ 上传失败: ' + e.message);
            }

            event.target.value = '';
        }

        async function webSearch() {
            const query = prompt('输入Web搜索关键词:');
            if (!query) return;

            addMessage('user', `Web搜索: ${query}`);
            addMessage('system', '⏳ 正在搜索网络信息（可能需要30-120秒）...');

            console.log('[webSearch] 开始Web搜索:', query);

            const loadingMsg = addLoadingMessage();

            try {
                const model = document.getElementById('model-select').value;
                console.log('[webSearch] 使用模型:', model);

                const response = await fetch('/api/web-search', {
                    method: 'POST',
                    headers: { 'Content-Type': 'application/json' },
                    body: JSON.stringify({
                        query,
                        model
                    })
                });

                console.log('[webSearch] 响应状态:', response.status);
                const data = await response.json();
                console.log('[webSearch] 响应数据:', Object.keys(data));

                // 如果是异步响应（202），需要轮询结果
                if (response.status === 202 && data.request_id) {
                    console.log('[webSearch] 异步任务ID:', data.request_id);
                    removeMessage(loadingMsg);
                    await pollWebSearchResult(data.request_id, query);
                } else if (data.results !== undefined && data.results !== null) {
                    console.log('[webSearch] 显示结果，长度:', data.results.length);
                    removeMessage(loadingMsg);
                    addMessage('assistant', data.results);
                } else if (data.error) {
                    console.log('[webSearch] 错误:', data.error);
                    removeMessage(loadingMsg);
                    addMessage('system', '❌ 搜索失败: ' + data.error);
                } else {
                    console.log('[webSearch] 无有效响应');
                    removeMessage(loadingMsg);
                    addMessage('system', '❌ 搜索失败: 无效的响应格式');
                }
            } catch (e) {
                console.log('[webSearch] 异常:', e.name, e.message);
                removeMessage(loadingMsg);
                addMessage('system', '❌ 搜索失败: ' + e.message);
            }
        }

        async function pollWebSearchResult(requestId, query) {
            // 轮询Web搜索结果
            const maxAttempts = 120;  // 120秒超时
            let attempts = 0;
            const loadingDiv = document.querySelector('.message.loading');

            while (attempts < maxAttempts) {
                try {
                    const response = await fetch(`/api/result/${requestId}`);
                    const data = await response.json();

                    console.log(`[webSearch] 轮询结果 [${attempts}]: status=${data.status}`);

                    if (data.status === 'completed') {
                        console.log('[webSearch] 搜索完成，开始处理结果...');

                        // 移除加载动画
                        if (loadingDiv && loadingDiv.parentNode) {
                            loadingDiv.parentNode.removeChild(loadingDiv);
                        }

                        // data.result包含_handle_web_search返回的{"results": ...., "query": ...}
                        if (data.result && data.result.results) {
                            console.log('[webSearch] 显示搜索结果');
                            addMessage('assistant', data.result.results);
                        } else {
                            console.warn('[webSearch] 结果为空，data.result=', data.result);
                            addMessage('system', '⚠️ 搜索完成但未获取结果');
                        }
                        console.log('[webSearch] 搜索结果已添加到消息');
                        return;
                    } else if (data.status === 'failed') {
                        console.log('[webSearch] 搜索失败:', data.error);

                        // 移除加载动画
                        if (loadingDiv && loadingDiv.parentNode) {
                            loadingDiv.parentNode.removeChild(loadingDiv);
                        }

                        addMessage('system', '❌ 搜索失败: ' + (data.error || '未知错误'));
                        return;
                    }

                    // status === 'running'，继续轮询
                    await new Promise(resolve => setTimeout(resolve, 1000));
                    attempts++;
                } catch (e) {
                    console.error('[webSearch] 轮询异常:', e);
                    await new Promise(resolve => setTimeout(resolve, 1000));
                    attempts++;
                }
            }

            // 移除加载动画
            if (loadingDiv && loadingDiv.parentNode) {
                loadingDiv.parentNode.removeChild(loadingDiv);
            }

            addMessage('system', '❌ 搜索超时（超过120秒）');
        }

        async function sendMessage() {
            const input = document.getElementById('input');
            const message = input.value.trim();

            if (!message || !sessionId) {
                if (!sessionId) addMessage('system', '❌ 会话未初始化');
                return;
            }

            input.value = '';
            addMessage('user', message);

            // 获取相关的对话历史
            const history = getRelevantHistory(message);

            // 显示加载动画
            const loadingMsg = addLoadingMessage();

            //前端发送请求传模型
            try {
                const model = document.getElementById('model-select').value;
                const kb = document.getElementById('kb-select').value;

                console.log('发送请求:', { message: message.substring(0, 50), model, kb, historyLength: history.length });

                const response = await fetch('/api/claude', {
                    method: 'POST',
                    headers: { 'Content-Type': 'application/json' },
                    body: JSON.stringify({
                        prompt: message,
                        model,
                        kb,
                        history: history  // 附加对话历史
                    })
                });

                console.log('初始响应状态:', response.status);

                if (!response.ok && response.status !== 202) {
                    removeMessage(loadingMsg);
                    throw new Error(`HTTP ${response.status}: ${response.statusText}`);
                }

                const data = await response.json();
                console.log('响应数据:', { status: response.status, data_keys: Object.keys(data) });

                // 检查是否是异步响应 (202)
                if (response.status === 202 && data.request_id) {
                    console.log('检测到异步响应，request_id:', data.request_id);

                    // 保存当前请求ID并显示停止按钮
                    currentRequestId = data.request_id;
                    shouldStopPolling = false;
                    document.getElementById('stop-btn').style.display = 'inline-block';

                    // 根据请求内容计算额外超时时间
                    let extraTimeoutMinutes = 0;
                    const msgLower = message.toLowerCase();
                    if (msgLower.includes('word') || msgLower.includes('docx') || msgLower.includes('文档')) {
                        extraTimeoutMinutes += 5;  // docx +5分钟
                        console.log('[Timeout] 检测到docx请求，增加5分钟超时');
                    }
                    if (msgLower.includes('ppt') || msgLower.includes('演示') || msgLower.includes('幻灯')) {
                        extraTimeoutMinutes += 5;  // pptx +5分钟
                        console.log('[Timeout] 检测到pptx请求，增加5分钟超时');
                    }

                    // 异步处理：轮询结果
                    await pollResult(data.request_id, loadingMsg, extraTimeoutMinutes);
                } else if (data.response) {
                    console.log('同步响应');
                    // 同步响应：直接显示
                    removeMessage(loadingMsg);
                    addMessage('assistant', data.response);
                } else if (data.result && data.result.response) {
                    console.log('异步完成响应');
                    // 异步完成：显示结果
                    removeMessage(loadingMsg);
                    addMessage('assistant', data.result.response);
                } else {
                    console.warn('无法找到响应内容，完整数据:', data);
                    removeMessage(loadingMsg);
                    addMessage('system', '⚠️ 无法解析服务器响应');
                }
            } catch (e) {
                console.error('sendMessage 异常:', e);
                removeMessage(loadingMsg);
                addMessage('system', '❌ 错误: ' + e.message);
            }
        }

        function addLoadingMessage() {
            const messagesDiv = document.getElementById('messages');
            const div = document.createElement('div');
            div.className = 'message loading';
            div.innerHTML = '<div class="spinner"></div><span>处理中...</span>';
            messagesDiv.appendChild(div);
            messagesDiv.scrollTop = messagesDiv.scrollHeight;
            return div;
        }

        function removeMessage(msgElement) {
            if (msgElement && msgElement.parentNode) {
                msgElement.parentNode.removeChild(msgElement);
            }
        }

        async function pollResult(requestId, loadingMsg, extraTimeoutMinutes = 0) {
            // 轮询异步请求的结果
            // 基础 5 分钟 + 额外时间（如docx +5分钟, pptx +5分钟）
            const baseAttempts = 300;  // 基础5分钟（每次1秒）
            const extraAttempts = extraTimeoutMinutes * 60;  // 额外分钟数转秒
            const maxAttempts = baseAttempts + extraAttempts;
            console.log(`[pollResult] 总超时时间: ${maxAttempts}秒 (基础300 + 额外${extraAttempts})`);
            let attempts = 0;

            while (attempts < maxAttempts) {
                // 检查是否需要停止轮询
                if (shouldStopPolling) {
                    console.log('[pollResult] 用户停止了轮询');
                    removeMessage(loadingMsg);
                    addMessage('system', '⏹️ 已停止进程');
                    document.getElementById('stop-btn').style.display = 'none';
                    shouldStopPolling = false;
                    currentRequestId = null;
                    return;
                }

                try {
                    const response = await fetch(`/api/result/${requestId}`);

                    if (!response.ok) {
                        console.warn(`轮询返回非200状态: ${response.status}`);
                        // 等待后重试
                        await new Promise(resolve => setTimeout(resolve, 1000));
                        attempts++;
                        continue;
                    }

                    const data = await response.json();
                    console.log(`轮询结果 [${attempts}]: status=${data.status}`);

                    if (data.status === 'completed') {
                        // 移除加载动画
                        removeMessage(loadingMsg);

                        // 显示最终结果
                        if (data.result) {
                            // 从嵌套的结果中提取响应
                            let responseText = null;
                            if (data.result.response) {
                                responseText = data.result.response;
                            } else if (typeof data.result === 'string') {
                                responseText = data.result;
                            }

                            if (responseText) {
                                addMessage('assistant', responseText);
                            } else {
                                console.log('完整结果对象:', data.result);
                                addMessage('system', '⚠️ 无法解析响应内容');
                            }
                        } else {
                            addMessage('system', '⚠️ 未获取到响应内容');
                        }
                        // 隐藏停止按钮
                        document.getElementById('stop-btn').style.display = 'none';
                        currentRequestId = null;
                        return;
                    } else if (data.status === 'failed') {
                        // 移除加载动画
                        removeMessage(loadingMsg);
                        addMessage('system', '❌ 处理失败: ' + (data.error || '未知错误'));
                        // 隐藏停止按钮
                        document.getElementById('stop-btn').style.display = 'none';
                        currentRequestId = null;
                        return;
                    }
                    // status === 'running'，继续轮询
                } catch (e) {
                    console.error(`轮询异常 [${attempts}]:`, e);
                    // 异常不应该停止轮询，继续重试
                }

                // 等待1秒后重试
                await new Promise(resolve => setTimeout(resolve, 1000));
                attempts++;
            }

            // 超时：移除加载动画
            removeMessage(loadingMsg);
            const totalMinutes = Math.ceil(maxAttempts / 60);
            addMessage('system', `❌ 请求超时（超过${totalMinutes}分钟）`);
            // 隐藏停止按钮
            document.getElementById('stop-btn').style.display = 'none';
            currentRequestId = null;
        }

        function clearMessages() {
            document.getElementById('messages').innerHTML = '';
            clearHistory();  // 同时清空对话历史
            addMessage('system', '✅ 聊天已清空，对话历史已重置');
        }

        async function stopProcess() {
            if (!currentRequestId) {
                addMessage('system', '⚠️ 当前没有运行的进程');
                return;
            }

            console.log('[stopProcess] 停止进程，request_id:', currentRequestId);
            addMessage('system', '⏳ 正在停止进程...');

            try {
                const response = await fetch('/api/stop', {
                    method: 'POST',
                    headers: {'Content-Type': 'application/json'},
                    body: JSON.stringify({request_id: currentRequestId})
                });

                if (response.ok) {
                    const result = await response.json();
                    shouldStopPolling = true;
                    addMessage('system', '✅ ' + (result.message || '进程已停止'));
                    console.log('[stopProcess] 停止成功:', result);
                } else {
                    addMessage('system', '⚠️ 停止请求失败: HTTP ' + response.status);
                }
            } catch (e) {
                console.error('[stopProcess] 错误:', e);
                addMessage('system', '❌ 停止失败: ' + e.message);
            }

            document.getElementById('stop-btn').style.display = 'none';
            currentRequestId = null;
        }

        function removeThinkTags(text) {
            // 安全检查：如果text为undefined、null或不是字符串，返回空字符串
            if (!text || typeof text !== 'string') {
                return '';
            }
            // 移除 <think>...</think> 标签及其内容
            var thinkStart = text.indexOf('<think>');
            var result = text;
            while (thinkStart >= 0) {
                var thinkEnd = result.indexOf('</think>', thinkStart);
                if (thinkEnd >= 0) {
                    result = result.substring(0, thinkStart) + result.substring(thinkEnd + 8);
                    thinkStart = result.indexOf('<think>');
                } else {
                    break;
                }
            }
            return result.trim();
        }

        function linkifyUrls(text) {
            // 安全检查：如果text为undefined、null或不是字符串，返回空字符串
            if (!text || typeof text !== 'string') {
                return '';
            }
            // 先转义HTML特殊字符
            const escaped = escapeHtml(text);

            // 使用正则表达式一次性替换所有URL
            const urlRegex = /https?:\/\/[^\s<>"'\)]+/g;
            const result = escaped.replace(urlRegex, function(url) {
                // 移除末尾的标点符号
                url = url.replace(/[.,;:!?\)]+$/, '');
                var displayUrl = url.length > 50 ? url.substring(0, 50) + '...' : url;
                return '<a href="' + url + '" target="_blank" rel="noopener noreferrer" style="color: #4a9eff; text-decoration: underline;">' + displayUrl + '</a>';
            });

            return result;
        }

        function downloadFile(filepath) {
            // 获取文件名
            const filename = filepath.split('/').pop();

            // 判断是否为完整路径
            const isFullPath = filepath.includes('/home/') || filepath.includes('/');

            // 构建下载URL
            let url;
            if (isFullPath && filepath.includes('/home/will/Downloads/')) {
                // 新生成的文件：直接用文件名搜索，后端会在下载目录找到
                url = `/api/download/${encodeURIComponent(filename)}`;
                console.log(`[Download] 检测为下载目录文件: ${filename}`);
            } else if (isFullPath) {
                // 其他路径：直接用文件名搜索
                url = `/api/download/${encodeURIComponent(filename)}`;
                console.log(`[Download] 检测为路径: ${filepath} → ${filename}`);
            } else {
                // 仅文件名
                url = `/api/download/${encodeURIComponent(filename)}`;
            }

            // 创建临时链接并触发下载
            const a = document.createElement('a');
            a.href = url;
            a.download = filename;
            document.body.appendChild(a);

            // 添加错误处理
            setTimeout(() => {
                a.click();
                setTimeout(() => {
                    document.body.removeChild(a);
                    addMessage('system', `⬇️ 已开始下载: ${filename}`);
                }, 100);
            }, 50);
        }

        function downloadKBFile(filename) {
            // 下载知识库文件
            // 先从当前KB目录尝试，再从所有KB目录搜索
            const currentKb = document.getElementById('kb-select').value || 'KB';
            const url = `/api/download/${currentKb}/${encodeURIComponent(filename)}`;

            const a = document.createElement('a');
            a.href = url;
            a.download = filename;
            document.body.appendChild(a);

            // 添加错误处理
            a.onclick = function(e) {
                fetch(url, {method: 'HEAD'})
                    .then(response => {
                        if (!response.ok) {
                            // 尝试直接文件名搜索（在所有KB中）
                            window.location.href = `/api/download/${encodeURIComponent(filename)}`;
                        } else {
                            a.click();
                        }
                    })
                    .catch(() => {
                        // 网络错误，尝试直接搜索
                        window.location.href = `/api/download/${encodeURIComponent(filename)}`;
                    });
                e.preventDefault();
                return false;
            };

            a.click();
            document.body.removeChild(a);

            addMessage('system', `📥 正在下载: ${filename}`);
        }

        function addMessage(role, content) {
            const messagesDiv = document.getElementById('messages');
            const div = document.createElement('div');
            div.className = 'message ' + role;

            // 格式化内容：保留换行和缩进
            if (role === 'assistant') {
                // 移除<think>...</think>标签
                content = removeThinkTags(content);

                // 如果内容为空（全是<think>标签），不显示
                if (!content) {
                    return;
                }

                // 使用pre标签保留格式，加上word-wrap样式，并识别URL
                const linkedContent = linkifyUrls(content);
                div.innerHTML = `<pre style="margin: 0; font-family: inherit; color: inherit; white-space: pre-wrap; word-break: break-word; overflow-wrap: break-word;">${linkedContent}</pre>`;
            } else {
                div.textContent = content;
            }

            messagesDiv.appendChild(div);
            messagesDiv.scrollTop = messagesDiv.scrollHeight;

            if (role === 'assistant') {
                addDownloadButtons(div, content);
            }

            // 记录对话历史（user和assistant角色）
            addToHistory(role, content);
        }

        function addDownloadButtons(messageDiv, content) {
            if (!content) return;

            const toolbar = document.createElement('div');
            toolbar.className = 'message-toolbar';

            const mdMatch = content.match(/\[点击下载: ([^\]]+\.md)\]\((http[^\)]+)\)/);
            const docxMatch = content.match(/\[Word文档: ([^\]]+\.docx)\]\((http[^\)]+)\)/);
            const pptxMatch = content.match(/\[PowerPoint演示文稿: ([^\]]+\.pptx)\]\((http[^\)]+)\)/);

            let hasButtons = false;

            if (mdMatch) {
                const btn = createDownloadButton('📄 Markdown', mdMatch[2], mdMatch[1]);
                toolbar.appendChild(btn);
                hasButtons = true;
            }
            if (docxMatch) {
                const btn = createDownloadButton('📝 Word', docxMatch[2], docxMatch[1]);
                toolbar.appendChild(btn);
                hasButtons = true;
            }
            if (pptxMatch) {
                const btn = createDownloadButton('📊 PPT', pptxMatch[2], pptxMatch[1]);
                toolbar.appendChild(btn);
                hasButtons = true;
            }

            if (hasButtons) {
                messageDiv.appendChild(toolbar);
            }
        }

        function createDownloadButton(label, url, filename) {
            const btn = document.createElement('button');
            btn.className = 'download-btn';
            btn.textContent = label;
            btn.onclick = (e) => {
                e.preventDefault();
                const a = document.createElement('a');
                a.href = url;
                a.download = filename;
                a.target = '_blank';
                document.body.appendChild(a);
                a.click();
                document.body.removeChild(a);
                addMessage('system', `⬇️ 正在下载: ${filename}`);
            };
            return btn;
        }

        function escapeHtml(text) {
            const div = document.createElement('div');
            div.textContent = text;
            return div.innerHTML;
        }

        document.addEventListener('DOMContentLoaded', () => {
            initSession();
            document.getElementById('input').addEventListener('keydown', (e) => {
                if (e.key === 'Enter' && e.shiftKey) {
                    e.preventDefault();
                    sendMessage();
                }
            });
        });
    </script>
</body>
</html>
        """
#  修改总结                                                                                                                                        
                                                                                                                                                  
#   1. 前端 - 对话历史管理 (server.py)                                                                                                              
                                                                                                                                                  
#   新增函数 (第1567-1655行):                                                                                                                       
#   - conversationHistory: 存储最近对话的数组                                                                                                       
#   - MAX_HISTORY_ROUNDS = 3: 最多保留3轮对话                                                                                                       
#   - isNewTopic(message): 判断是否是新话题（搜索、查找、请问等）                                                                                   
#   - isContextualRequest(message): 判断是否引用上文（上面、转换、生成等）                                                                          
#   - getRelevantHistory(currentMessage): 智能获取相关历史                                                                                          
#   - addToHistory(role, content): 添加消息到历史                                                                                                   
#   - clearHistory(): 清空历史                                                                                                                      
                                                                                                                                                  
#   修改 sendMessage (第1741-1774行):                                                                                                               
#   - 获取相关对话历史                                                                                                                              
#   - 发送时附加 history 参数                                                                                                                       
                                                                                                                                                  
#   修改 addMessage (第2042-2073行):                                                                                                                
#   - 调用 addToHistory 记录对话                                                                                                                    
                                                                                                                                                  
#   修改 clearMessages (第1914-1918行):                                                                                                             
#   - 同时清空对话历史                                                                                                                              
                                                                                                                                                  
#   2. 后端 - 接收和传递历史 (server.py)                                                                                                            
                                                                                                                                                  
#   修改 _handle_claude_call:                                                                                                     
#   - 获取 history 参数                                                                                                                             
#   - 记录历史条数日志                                                                                                                              
#   - 传递历史给语义调度系统                                                                                                                        
                                                                                                                                                  
#   3. 调度系统 - 使用对话历史 (scheduler.py)                                                                                                       
                                                                                                                                                  
#   修改 process 方法 (第1908-1958行):                                                                                                              
#   - 从 context 获取 history                                                                                                                       
#   - 格式化历史为文本                                                                                                                              
#   - 构建带上下文的输入                                                                                                                            
#   - 仅对知识库和通用处理使用上下文                                                                                                                
                                                                                                                                                  
#   新增 _format_history 方法 (第1960-1975行):                                                                                                      
#   - 格式化对话历史为可读文本                                                                                                                      
#   - 截断过长内容（500字符）                                                                                                                       
                                                                                                                                                  
#   工作流程示例                                                                                                                                    
                                                                                                                                                  
#   用户: "总结道德经核心思想"                                                                                                                      
#     → 执行知识库搜索，返回总结                                                                                                                    
#     → 历史记录: [{role:'user', content:'总结道德经...'}, {role:'assistant', content:'...总结内容...'}]                                            
                                                                                                                                                  
#   用户: "把上面内容转换为Word文档"                                                                                                                
#     → 检测到引用上文（"上面"、"Word"）                                                                                                            
#     → 附加历史上下文                                                                                                                              
#     → 使用历史中的总结内容生成Word文档                                                                                                            
                                                                                                                                                  
#   用户: "搜索论语"                                                                                                                                
#     → 检测到新话题（"搜索"）                                                                                                                      
#     → 不附加历史，开始新对话   
# ============ 启动服务器 ============

def cleanup_sessions():
    """清理过期会话和请求记录"""
    while True:
        time.sleep(10)  # 改为10秒运行一次，更频繁地清理
        with SESSIONS_LOCK:
            now = datetime.now()
            # 清理过期的会话
            expired = [
                sid for sid, s in SESSIONS.items()
                if (now - s.last_activity).seconds > SESSION_TIMEOUT
            ]
            for sid in expired:
                del SESSIONS[sid]
                logger.info(f"清理过期会话: {sid}")

            # 清理已完成的请求记录（保留最多500条）
            if len(PROGRESS_TRACKING) > 500:
                completed_requests = [
                    req_id for req_id, data in PROGRESS_TRACKING.items()
                    if data.get("status") in ["completed", "error"]
                ]
                # 删除最早的已完成记录，保留最新的
                for req_id in completed_requests[:-400]:  # 保留400条
                    del PROGRESS_TRACKING[req_id]
                logger.info(f"清理请求记录: 删除{len(completed_requests) - 400}条")

if __name__ == "__main__":
    logger.info(f"启动Claude Web Server LLM在端口 {WEB_PORT}...")
    logger.info(f"访问地址: http://localhost:{WEB_PORT}")

    # 初始化语义调度系统

    #SemanticScheduler-从SemanticAnalyzer继承
    #SemanticScheduler-Dispatcher
    # analyzer =
    # <scheduler.SemanticAnalyzer object at 0x724e4faf8f50>
    # discovery =
    # <scheduler.SkillDiscovery object at 0x724e4c2e3150>
    # dispatcher =
    # <scheduler.Dispatcher object at 0x724e4fafa150>
    try:
        semantic_scheduler = SemanticScheduler(SKILLS_DIR, scheduler_model_caller, KB_PATHS)
        skills_count = len(semantic_scheduler.discovery.skills)
        agents_count = len(semantic_scheduler.discovery.agents)
        logger.info(f"✅ 语义调度系统初始化成功")
        logger.info(f"   - 发现 {skills_count} 个 Skills")
        logger.info(f"   - 可用 {agents_count} 个 Agents")
        logger.info(f"   - Skills 目录: {SKILLS_DIR}")
    except Exception as e:
        logger.warning(f"⚠️ 语义调度系统初始化失败: {e}")
        semantic_scheduler = None

    # 启动清理线程
    cleanup_thread = threading.Thread(target=cleanup_sessions, daemon=True)
    cleanup_thread.start()

    # 启动HTTP服务器
    class ReuseAddrTCPServer(socketserver.TCPServer):
        allow_reuse_address = True

    with ReuseAddrTCPServer(("0.0.0.0", WEB_PORT), WebHandler) as httpd:
        try:
            httpd.serve_forever()
        except KeyboardInterrupt:
            logger.info("\n服务器已停止")
